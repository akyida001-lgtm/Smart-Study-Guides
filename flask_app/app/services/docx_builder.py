"""Build a .docx file from the formatted assignment text."""
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Headings that signal a new major section
_SECTION_HEADINGS = {
    "INTRODUCTION", "INTRODUCTION PARAGRAPH",
    "MAIN CONTENT", "MAIN CONTENT/BODY SECTION", "BODY", "BODY PARAGRAPHS",
    "CONCLUSION", "CONCLUSION PARAGRAPH",
}

# Headings that trigger a page break + hanging-indent reference formatting
_REF_HEADINGS = {
    "REFERENCES", "WORKS CITED", "BIBLIOGRAPHY",
    "REFERENCE LIST", "WORKS CITED PAGE",
}


def _set_font(run, bold=False, size=12):
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def _normal_para(doc, text, first_line_indent=True):
    """Body paragraph: double-spaced, 0.5-inch first-line indent."""
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.line_spacing = 2.0
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_line_indent:
        fmt.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def _hanging_para(doc, text):
    """Reference entry: double-spaced, 0.5-inch hanging indent."""
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.line_spacing = 2.0
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.left_indent = Inches(0.5)
    fmt.first_line_indent = Inches(-0.5)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def _add_centered(doc, text, bold=False, size=12, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    return p


def _build_title_page(doc, topic, course_name, student_name,
                      instructor_name, school_name, due_date):
    for _ in range(4):
        doc.add_paragraph("")
    _add_centered(doc, topic.strip(), bold=True, size=12)
    for _ in range(3):
        doc.add_paragraph("")
    if student_name:
        _add_centered(doc, student_name)
    if school_name:
        _add_centered(doc, school_name)
    if course_name:
        _add_centered(doc, course_name)
    if instructor_name:
        _add_centered(doc, instructor_name)
    if due_date:
        _add_centered(doc, due_date)
    doc.add_page_break()


def build_docx(text: str, topic: str, *,
               course_name: str = "", student_name: str = "",
               instructor_name: str = "", school_name: str = "",
               due_date: str = "") -> bytes:
    doc = Document()

    # Global default style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # 1-inch margins
    section = doc.sections[0]
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    _build_title_page(doc, topic,
                      course_name=course_name,
                      student_name=student_name,
                      instructor_name=instructor_name,
                      school_name=school_name,
                      due_date=due_date)

    # Strip markdown code fences if AI wrapped the output
    text = text.strip()
    if text.startswith("```"):
        text = text.split("\n", 1)[-1]
    if text.endswith("```"):
        text = text.rsplit("\n", 1)[0]
    text = text.strip()

    in_title_section = False
    in_references    = False
    ref_page_added   = False

    # Add topic title at the top of the body (before Introduction)
    _add_centered(doc, topic.strip(), bold=True, size=12)

    for raw in text.split("\n"):
        line     = raw.rstrip()
        stripped = line.strip()
        upper    = stripped.upper()

        # Skip AI-generated title page block
        if upper == "TITLE PAGE":
            in_title_section = True
            continue
        if in_title_section:
            if upper in _SECTION_HEADINGS or upper in _REF_HEADINGS or stripped.startswith("#"):
                in_title_section = False
            else:
                continue

        if not stripped:
            doc.add_paragraph("")
            continue

        # Detect references section
        bare = stripped.lstrip("# ").strip().upper()
        if bare in _REF_HEADINGS:
            if not ref_page_added:
                doc.add_page_break()
                ref_page_added = True
            h = doc.add_heading(stripped.lstrip("# ").strip(), level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in h.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.font.bold = True
            in_references = True
            continue

        # Reference entries (hanging indent)
        if in_references:
            _hanging_para(doc, stripped)
            continue

        # Subheadings (## …)
        if stripped.startswith("##"):
            subheading_text = stripped.lstrip("# ").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            run = p.add_run(subheading_text)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            continue

        # Main section headings (# … or keyword)
        if stripped.startswith("#") or upper in _SECTION_HEADINGS:
            heading_text = stripped.lstrip("# ").strip()
            h = doc.add_heading(heading_text, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in h.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.font.bold = True
            continue

        # Regular body paragraph
        _normal_para(doc, stripped, first_line_indent=True)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
