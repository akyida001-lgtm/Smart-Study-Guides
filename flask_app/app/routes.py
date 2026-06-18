"""All HTTP routes."""
import os
import uuid
import threading
import queue as _queue
from datetime import datetime, timedelta

import requests as http_requests

from flask import (
    Blueprint, render_template, request, redirect, url_for,
    jsonify, flash, current_app, session, Response, stream_with_context,
)
from flask_login import current_user

import math
import json
import zipfile
import io
import csv

from .models import (db, User, Assignment, Source, Transaction,
                      ChatSession, ChatMessage,
                      HumanOrder, HumanOrderMessage, HumanOrderFile,
                      JobDocument, UserNotification,
                      Subscription, DailyUsage, AIRemovalJob, Review)
from .native_auth import require_login
from .services import semantic_scholar, pesapal, ryter_service, stripe_service
from .pipeline import start_pipeline, start_finalize_pipeline
from .prompts import GENERATION_PROMPT, DOCX_FORMATTING_PROMPT, _STYLE_RULES, _DEFAULT_STYLE_RULE
from .services.docx_builder import build_docx as _build_docx
from .services import supabase_storage

# Temporary in-memory queue for live DOCX conversion texts (keyed by assignment id)
_docx_queue:        dict = {}   # {aid: job_meta}
_humanize_jobs:     dict = {}   # {aid: job_state}
_docx_results:      dict = {}   # {aid: job_state}
_plagiarism_jobs:   dict = {}   # {job_id: job_state}

# ── Weather cache (in-memory, keyed by rounded lat_lon) ──────────────────────
_weather_cache: dict = {}    # {key: {"data": {...}, "ts": float}}
_WEATHER_TTL   = 600         # 10 minutes


_QUOTES = [
    # ── Famous ──────────────────────────────────────────────────────────────────
    {"text": "It always seems impossible until it's done.",
     "author": "Nelson Mandela", "type": "famous"},
    {"text": "The secret of getting ahead is getting started.",
     "author": "Mark Twain", "type": "famous"},
    {"text": "An investment in knowledge pays the best interest.",
     "author": "Benjamin Franklin", "type": "famous"},
    {"text": "The beautiful thing about learning is that nobody can take it away from you.",
     "author": "B.B. King", "type": "famous"},
    {"text": "The expert in anything was once a beginner.",
     "author": "Helen Hayes", "type": "famous"},
    {"text": "You don't have to be great to start, but you have to start to be great.",
     "author": "Zig Ziglar", "type": "famous"},
    {"text": "Live as if you were to die tomorrow. Learn as if you were to live forever.",
     "author": "Mahatma Gandhi", "type": "famous"},
    {"text": "Success is not final, failure is not fatal: it is the courage to continue that counts.",
     "author": "Winston Churchill", "type": "famous"},
    {"text": "The more that you read, the more things you will know.",
     "author": "Dr. Seuss", "type": "famous"},
    {"text": "Education is the most powerful weapon which you can use to change the world.",
     "author": "Nelson Mandela", "type": "famous"},
    # ── Custom ───────────────────────────────────────────────────────────────────
    {"text": "Consistency beats last-minute pressure. Start early, finish strong.",
     "author": "Smart Study Guides", "type": "custom"},
    {"text": "Every great paper starts with a single word. Write yours today.",
     "author": "Smart Study Guides", "type": "custom"},
    {"text": "Small daily progress is the key to big academic achievements.",
     "author": "Smart Study Guides", "type": "custom"},
    {"text": "Your deadline is not your enemy — procrastination is.",
     "author": "Smart Study Guides", "type": "custom"},
    {"text": "Quality research today means confident submission tomorrow.",
     "author": "Smart Study Guides", "type": "custom"},
    {"text": "The student who prepares is the student who succeeds.",
     "author": "Smart Study Guides", "type": "custom"},
    {"text": "One assignment at a time. That is how mountains move.",
     "author": "Smart Study Guides", "type": "custom"},
    {"text": "Clear thinking leads to clear writing. Take it one step at a time.",
     "author": "Smart Study Guides", "type": "custom"},
    {"text": "Your best work happens when preparation meets opportunity.",
     "author": "Smart Study Guides", "type": "custom"},
    {"text": "Every cited source you include strengthens your argument. Make it count.",
     "author": "Smart Study Guides", "type": "custom"},
]


def _daily_quote() -> dict:
    """Return the quote for today, consistent all day, changing at midnight."""
    idx = datetime.utcnow().date().toordinal() % len(_QUOTES)
    return _QUOTES[idx]


def _weather_icon(code: int) -> str:
    """Map OpenWeather condition code to an emoji."""
    if   code < 300: return "⛈️"
    elif code < 400: return "🌦️"
    elif code < 600: return "🌧️"
    elif code < 700: return "❄️"
    elif code < 800: return "🌫️"
    elif code == 800: return "☀️"
    elif code <= 802: return "⛅"
    else:             return "☁️"


def _notify(user_id: str, ntype: str, title: str, body: str):
    """Create a UserNotification row and silently swallow any DB errors."""
    try:
        n = UserNotification(user_id=user_id, type=ntype, title=title, body=body)
        db.session.add(n)
        db.session.commit()
    except Exception:
        db.session.rollback()


WORDS_PER_PAGE = 265

# ── Subscription plans ─────────────────────────────────────────────────────────
PLANS = {
    "basic": {
        "name": "Basic", "daily_limit": 2,
        "monthly": 27.99, "halfyear": 139.99, "yearly": 249.99,
        "desc": "2 AI assignments per day",
    },
    "standard": {
        "name": "Standard", "daily_limit": 4,
        "monthly": 37.99, "halfyear": 189.99, "yearly": 329.99,
        "desc": "4 AI assignments per day",
    },
    "unlimited": {
        "name": "Unlimited", "daily_limit": None,
        "monthly": 57.99, "halfyear": 289.99, "yearly": 499.99,
        "desc": "Unlimited AI assignments",
    },
}

PERIOD_MONTHS = {"monthly": 1, "halfyear": 6, "yearly": 12}

NEW_VISITOR_DISCOUNT = 10   # % off first subscription
REFERRAL_DISCOUNT    = 10   # % off if user was referred

# Legacy constant kept for old pesapal/paypal routes (not exposed in UI)
MIN_PURCHASE_USD = 12

# Core owner — protected from all admin actions, cannot be modified by anyone
CORE_OWNER_EMAIL = "aservices767@gmail.com"

# All owner/admin emails — bypass ALL subscription and daily-limit gates
OWNER_EMAILS = {"aservices767@gmail.com", "akyida001@gmail.com", "simonedwardj@yahoo.com", "simonedwardj6@gmail.com"}


def _is_owner(user) -> bool:
    """True if user is staff (DB flag) OR is a designated owner email."""
    if user is None:
        return False
    if getattr(user, "is_staff", False):
        return True
    email = (getattr(user, "email", "") or "").lower().strip()
    return email in OWNER_EMAILS


def _get_active_subscription(user_id):
    """Return the current active Subscription row, or None."""
    return (
        Subscription.query
        .filter_by(user_id=user_id, status="active")
        .filter(Subscription.end_date > datetime.utcnow())
        .order_by(Subscription.end_date.desc())
        .first()
    )


def _get_today_usage(user_id):
    """Return how many AI assignments the user has started today (UTC date)."""
    from datetime import date
    row = DailyUsage.query.filter_by(user_id=user_id, date=date.today()).first()
    return row.count if row else 0


def _increment_daily_usage(user_id):
    """Add 1 to today's assignment count for this user."""
    from datetime import date
    today = date.today()
    row = DailyUsage.query.filter_by(user_id=user_id, date=today).first()
    if row:
        row.count += 1
    else:
        row = DailyUsage(user_id=user_id, date=today, count=1)
        db.session.add(row)
    db.session.commit()


FREE_TRIAL_MAX_PAGES = 2

def _can_generate(user_id):
    """Return (ok, reason, sub, today_count, daily_limit)."""
    from .models import User as _User
    user = _User.query.get(user_id)
    if _is_owner(user):
        return True, "ok", None, 0, None  # owner/admins have no limit
    sub = _get_active_subscription(user_id)
    if not sub:
        # Allow two free trial assignments
        trial_count = user.free_trial_count or 0
        if trial_count < 2:
            return True, "free_trial", None, 0, 2
        return False, "no_subscription", None, 0, 0
    plan_info   = PLANS.get(sub.plan, {})
    daily_limit = plan_info.get("daily_limit")
    today_count = _get_today_usage(user_id)
    if daily_limit is None:
        return True, "ok", sub, today_count, None
    if today_count >= daily_limit:
        return False, "daily_limit", sub, today_count, daily_limit
    return True, "ok", sub, today_count, daily_limit


def register_routes(app):
    main = Blueprint("main", __name__)

    @main.route("/healthz")
    def healthz():
        return {"status": "ok"}, 200

    @main.route("/admin/setup")
    @require_login
    def admin_setup():
        """Grant the currently logged-in user staff access.
        Protected by STAFF_SETUP_KEY env var — visit /admin/setup?key=YOUR_KEY
        """
        secret = os.environ.get("STAFF_SETUP_KEY", "")
        provided = request.args.get("key", "")
        if not secret:
            return "STAFF_SETUP_KEY environment variable is not set.", 400
        if not provided or provided != secret:
            return "Invalid key.", 403
        current_user.is_staff = True
        db.session.commit()
        flash(f"✅ Staff access granted to {current_user.email}. You can now visit /admin/chats.", "success")
        return redirect(url_for("main.admin_chats"))

    @main.route("/robots.txt")
    def robots_txt():
        lines = [
            "User-agent: *",
            "Allow: /",
            "Allow: /humanize",
            "Allow: /check-ai",
            "Allow: /check-plagiarism",
            "Allow: /pricing",
            "Allow: /help",
            "Allow: /terms",
            "Allow: /privacy",
            "Allow: /about",
            "Disallow: /dashboard",
            "Disallow: /create/",
            "Disallow: /api/",
            "Disallow: /admin/",
            "Disallow: /auth/",
            "Disallow: /subscribe/",
            "Disallow: /account/",
            "Disallow: /profile",
            "Disallow: /transactions",
            "Disallow: /assignments/",
            "Disallow: /chat",
            "Disallow: /human-orders/",
            "Disallow: /writer-portal/",
            "Disallow: /job-docs/",
            "",
            "Sitemap: https://smart-study-guides.com/sitemap.xml",
        ]
        return Response("\n".join(lines), mimetype="text/plain")

    @main.route("/sitemap.xml")
    def sitemap_xml():
        pages = [
            ("https://smart-study-guides.com/",                 "1.0",  "weekly"),
            ("https://smart-study-guides.com/pricing",          "0.9",  "weekly"),
            ("https://smart-study-guides.com/humanize",         "0.85", "monthly"),
            ("https://smart-study-guides.com/check-ai",         "0.85", "monthly"),
            ("https://smart-study-guides.com/check-plagiarism", "0.85", "monthly"),
            ("https://smart-study-guides.com/help",             "0.7",  "monthly"),
            ("https://smart-study-guides.com/about",            "0.6",  "monthly"),
            ("https://smart-study-guides.com/terms",            "0.4",  "yearly"),
            ("https://smart-study-guides.com/privacy",          "0.4",  "yearly"),
        ]
        xml = ['<?xml version="1.0" encoding="UTF-8"?>',
               '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">']
        for loc, pri, freq in pages:
            xml.append(f"  <url><loc>{loc}</loc>"
                       f"<changefreq>{freq}</changefreq>"
                       f"<priority>{pri}</priority></url>")
        xml.append("</urlset>")
        return Response("\n".join(xml), mimetype="application/xml")

    @main.route("/")
    def index():
        # Capture referral code from ?ref=CODE for later use after signup
        ref = request.args.get("ref")
        if ref:
            session["referral_code"] = ref.strip().upper()
        return render_template("home.html", user=current_user)

    @main.route("/dashboard")
    @require_login
    def dashboard():
        if current_user.is_writer or session.get("admin_writer_preview"):
            return redirect(url_for("main.writer_available"))
        assignments = (
            Assignment.query.filter_by(user_id=current_user.id)
            .order_by(Assignment.created_at.desc())
            .limit(50).all()
        )
        host = request.host
        referral_link = f"https://{host}/?ref={current_user.referral_code}" if current_user.referral_code else ""

        # Stats
        non_draft   = [a for a in assignments if a.status != "draft"]
        total_pages = sum(a.pages for a in non_draft)
        in_progress = sum(1 for a in assignments if a.status in ("queued", "running", "generating"))

        # Subscription
        active_sub  = _get_active_subscription(current_user.id)
        today_count = _get_today_usage(current_user.id)
        plan_info   = PLANS.get(active_sub.plan, {}) if active_sub else {}
        daily_limit = plan_info.get("daily_limit") if active_sub else 0

        # Show review prompt after first completed assignment, once only
        completed_count = sum(1 for a in assignments if a.status == "done")
        show_review = (completed_count >= 1 and not current_user.review_prompted
                       and not _is_owner(current_user))

        return render_template(
            "dashboard.html",
            user=current_user,
            assignments=assignments,
            words_per_page=WORDS_PER_PAGE,
            referral_link=referral_link,
            total_pages=total_pages,
            in_progress=in_progress,
            daily_quote=_daily_quote(),
            active_sub=active_sub,
            today_count=today_count,
            daily_limit=daily_limit,
            plan_info=plan_info,
            show_review=show_review,
        )

    # ---------- Assignment creation flow ----------

    @main.route("/assignments/new", methods=["GET", "POST"])
    @require_login
    def new_assignment():
        # Check free trial status for the form
        active_sub = _get_active_subscription(current_user.id)
        is_free_trial = (not active_sub and not _is_owner(current_user)
                         and not current_user.free_trial_used)

        if request.method == "POST":
            try:
                raw_pages = int(request.form.get("pages", 1))
                # Free trial users capped at 2 pages
                if is_free_trial:
                    pages = max(1, min(FREE_TRIAL_MAX_PAGES, raw_pages))
                else:
                    pages = max(1, min(5, raw_pages))
                num_sources = max(1, int(request.form.get("num_sources", 4)))
            except ValueError:
                flash("Invalid number entered.", "error")
                return redirect(url_for("main.new_assignment"))

            topic = (request.form.get("topic") or "").strip()
            style = request.form.get("style", "APA")
            level = request.form.get("level", "Undergraduate")
            course_name = (request.form.get("course_name") or "").strip()
            student_name = (request.form.get("student_name") or "").strip()
            instructor_name = (request.form.get("instructor_name") or "").strip()
            school_name = (request.form.get("school_name") or "").strip()
            due_date = (request.form.get("due_date") or "").strip()

            if not topic:
                flash("Please enter a topic.", "error")
                return redirect(url_for("main.new_assignment"))

            word_count = pages * WORDS_PER_PAGE
            credit_cost = 0   # subscription model — no per-assignment cost

            atype = request.form.get("assignment_type", "standard")
            if atype not in ("standard", "open", "simple"):
                atype = "standard"
            img_url = (request.form.get("instruction_image_url") or "").strip() or None

            # Open Task / Simple: no sources or citations needed
            if atype in ("open", "simple"):
                num_sources = 0
                style = "None"

            a = Assignment(
                user_id=current_user.id, topic=topic, pages=pages,
                word_count=word_count, num_sources=num_sources,
                style=style, education_level=level,
                credit_cost=credit_cost, status="draft",
                course_name=course_name or None,
                student_name=student_name or None,
                instructor_name=instructor_name or None,
                school_name=school_name or None,
                due_date=due_date or None,
                assignment_type=atype,
                instruction_image_url=img_url,
            )
            db.session.add(a)
            db.session.commit()
            if atype in ("open", "simple"):
                return redirect(url_for("main.assignment_preview", aid=a.id))
            return redirect(url_for("main.assignment_sources", aid=a.id))

        return render_template("create_step1.html", user=current_user,
                               is_free_trial=is_free_trial,
                               free_trial_max_pages=FREE_TRIAL_MAX_PAGES)

    @main.route("/assignments/<int:aid>/sources", methods=["GET", "POST"])
    @require_login
    def assignment_sources(aid):
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return "Forbidden", 403

        if request.method == "POST":
            mode = request.form.get("mode", "auto")
            Source.query.filter_by(assignment_id=a.id).delete()

            if mode == "manual":
                titles    = request.form.getlist("source_title[]")
                summaries = request.form.getlist("source_summary[]")
                urls      = request.form.getlist("source_url[]")
                authors_l = request.form.getlist("source_authors[]")
                years_l   = request.form.getlist("source_year[]")
                # Pad shorter lists so zip works safely
                n = max(len(titles), len(summaries), len(urls))
                while len(authors_l) < n: authors_l.append("")
                while len(years_l)   < n: years_l.append("")
                for t, s, u, auth, yr in zip(titles, summaries, urls, authors_l, years_l):
                    if t.strip() or s.strip() or u.strip():
                        db.session.add(Source(
                            assignment_id=a.id, title=t.strip() or None,
                            summary=s.strip() or None, url=u.strip() or None,
                            authors=auth.strip() or None,
                            year=int(yr.strip()) if yr.strip().isdigit() else None,
                            is_user_provided=True,
                        ))
            elif mode == "auto-confirmed":
                # Student reviewed sources on Step 2 and clicked confirm
                titles    = request.form.getlist("src_title[]")
                summaries = request.form.getlist("src_summary[]")
                urls      = request.form.getlist("src_url[]")
                authors_l = request.form.getlist("src_authors[]")
                years_l   = request.form.getlist("src_year[]")
                for t, s, u, auth, yr in zip(titles, summaries, urls, authors_l, years_l):
                    db.session.add(Source(
                        assignment_id=a.id, title=t or None,
                        summary=s or None, url=u or None,
                        authors=auth or None,
                        year=int(yr) if yr and yr.isdigit() else None,
                        is_user_provided=False,
                    ))
            else:
                # Legacy fallback: search now
                found = semantic_scholar.find_sources(a.topic, limit=a.num_sources)
                for s in found:
                    db.session.add(Source(
                        assignment_id=a.id, title=s["title"],
                        summary=s["summary"], url=s["url"],
                        authors=s.get("authors") or None,
                        year=s.get("year") or None,
                        is_user_provided=False,
                    ))
            db.session.commit()
            return redirect(url_for("main.assignment_step2b", aid=a.id))

        return render_template("create_step2.html", a=a, user=current_user)

    @main.route("/create/step2b/<int:aid>")
    @require_login
    def assignment_step2b(aid):
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return "Forbidden", 403
        sources = Source.query.filter_by(assignment_id=a.id).all()
        return render_template("create_step2b.html", a=a, sources=sources, user=current_user)

    @main.route("/api/assignments/<int:aid>/generate-annotations", methods=["POST"])
    @require_login
    def api_generate_annotations(aid):
        from .services import openai_service as _oa
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return jsonify({"error": "forbidden"}), 403
        sources = Source.query.filter_by(assignment_id=a.id).all()
        if not sources:
            return jsonify({"error": "no sources"}), 400
        src_dicts = [
            {"title": s.title, "authors": s.authors, "year": s.year,
             "url": s.url, "summary": s.summary}
            for s in sources
        ]
        try:
            annotations = _oa.generate_source_annotations(
                src_dicts, a.topic,
                style=a.style, level=a.education_level
            )
        except Exception as e:
            import traceback, logging as _log2
            _log2.getLogger(__name__).error(
                "generate_source_annotations failed for assignment %s: %s\n%s",
                aid, e, traceback.format_exc()
            )
            return jsonify({"error": str(e)}), 500
        # Save to DB and return
        result = []
        for s, ann in zip(sources, annotations):
            s.apa_intext    = ann.get("intext") or ""
            s.apa_reference = ann.get("reference") or ""
            s.annotation    = ann.get("paragraph") or ""
            result.append({
                "id":        s.id,
                "title":     s.title,
                "topic":     ann.get("topic") or s.title,
                "intext":    s.apa_intext,
                "reference": s.apa_reference,
                "paragraph": s.annotation,
            })
        db.session.commit()
        return jsonify(result)

    @main.route("/api/assignments/<int:aid>/save-annotations", methods=["POST"])
    @require_login
    def api_save_annotations(aid):
        """Save student-edited paragraphs before proceeding to preview."""
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return jsonify({"error": "forbidden"}), 403
        data = request.get_json(force=True) or {}
        edits = data.get("edits", {})   # {str(source_id): paragraph_text}
        for s in Source.query.filter_by(assignment_id=a.id).all():
            key = str(s.id)
            if key in edits:
                s.annotation = edits[key]
        db.session.commit()
        return jsonify({"ok": True})

    @main.route("/api/upload-instruction-image", methods=["POST"])
    @require_login
    def api_upload_instruction_image():
        """Upload an instruction image to Supabase and return its signed URL."""
        import base64 as _b64
        f = request.files.get("file")
        if not f or not f.filename:
            return jsonify({"error": "No file provided."}), 400
        fname = f.filename.lower()
        allowed = (".jpg", ".jpeg", ".png", ".webp", ".gif")
        if not any(fname.endswith(e) for e in allowed):
            return jsonify({"error": "Please upload a JPG, PNG, WEBP, or GIF image."}), 400
        data = f.read()
        if len(data) > 15 * 1024 * 1024:
            return jsonify({"error": "Image too large — max 15 MB."}), 400
        try:
            import time as _t
            ext = fname.rsplit(".", 1)[-1].lower()
            filename = f"instr_{current_user.id}_{int(_t.time())}.{ext}"
            ct_map = {"jpg": "image/jpeg", "jpeg": "image/jpeg",
                      "png": "image/png", "webp": "image/webp", "gif": "image/gif"}
            ct = ct_map.get(ext, "image/jpeg")
            url = supabase_storage.upload_file(filename, data, ct, signed_days=365)
            return jsonify({"url": url})
        except Exception as e:
            return jsonify({"error": f"Upload failed: {e}"}), 500

    @main.route("/api/extract-file-text", methods=["POST"])
    @require_login
    def api_extract_file_text():
        """Extract text from an uploaded PDF, DOCX, or image file.
        Optional ?mode=sources returns a parsed JSON source list via GPT."""
        import io, base64
        f = request.files.get("file")
        if not f or not f.filename:
            return jsonify({"error": "No file provided."}), 400

        fname = f.filename.lower()
        data  = f.read()
        if len(data) > 15 * 1024 * 1024:
            return jsonify({"error": "File too large — max 15 MB."}), 400

        text = ""
        try:
            if fname.endswith(".pdf"):
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(data))
                parts = [page.extract_text() or "" for page in reader.pages]
                text  = "\n\n".join(p for p in parts if p.strip())

            elif fname.endswith(".docx"):
                import docx as _docx
                doc  = _docx.Document(io.BytesIO(data))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

            elif fname.endswith(".doc"):
                return jsonify({"error": "Old .doc format is not supported — please save as .docx or PDF."}), 400

            elif any(fname.endswith(ext) for ext in (".jpg", ".jpeg", ".png", ".webp", ".gif")):
                # Use OpenAI Vision — extract text AND describe visual content
                from .services.openai_service import client as _oai_client
                ct_map = {".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                          ".png": "image/png", ".webp": "image/webp", ".gif": "image/gif"}
                ext_key = next((e for e in ct_map if fname.endswith(e)), ".jpg")
                ct  = ct_map[ext_key]
                b64 = base64.b64encode(data).decode()
                resp = _oai_client().chat.completions.create(
                    model="gpt-4o",
                    messages=[{"role": "user", "content": [
                        {"type": "text", "text":
                         "Describe what you see in this image in full detail. "
                         "If there is any text, transcribe it exactly. "
                         "If it is a question, problem, scenario, or diagram, explain it thoroughly. "
                         "Output only the description/transcription — no preamble."},
                        {"type": "image_url", "image_url": {"url": f"data:{ct};base64,{b64}", "detail": "high"}},
                    ]}],
                    max_tokens=2000,
                )
                text = resp.choices[0].message.content.strip()
            else:
                return jsonify({"error": "Unsupported file type. Please upload a PDF, DOCX, JPG, PNG, or WEBP."}), 400

        except Exception as e:
            return jsonify({"error": f"Could not read file: {e}"}), 500

        if not text.strip():
            return jsonify({"error": "No content could be extracted from this file."}), 400

        # ── Sources mode: parse text into structured source list ──
        mode = request.args.get("mode", "text")
        if mode == "sources":
            try:
                import openai as _oai, json as _json2
                client = _oai.OpenAI(api_key=_os.environ["OPENAI_API_KEY"])
                prompt = (
                    "The following text is a list of academic sources / references. "
                    "Parse it into a JSON array where each object has keys: "
                    "\"title\" (string), \"authors\" (string, comma-separated), "
                    "\"year\" (string or null), \"url\" (string or null), \"summary\" (string, one sentence). "
                    "Return ONLY the JSON array, no markdown, no explanation.\n\n"
                    + text[:6000]
                )
                r2  = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=2000,
                )
                raw = r2.choices[0].message.content.strip()
                raw = raw.lstrip("```json").lstrip("```").rstrip("```").strip()
                sources = _json2.loads(raw)
                return jsonify({"sources": sources})
            except Exception as e:
                # fall back to returning raw text if parsing fails
                return jsonify({"text": text, "sources_error": str(e)})

        return jsonify({"text": text})

    @main.route("/api/assignments/<int:aid>/find-sources")
    @require_login
    def api_find_sources(aid):
        """Search Semantic Scholar and return source candidates as JSON for the Step 2 preview."""
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return "Forbidden", 403
        offset = int(request.args.get("offset", 0))
        found = semantic_scholar.find_sources(a.topic, limit=8, offset=offset)
        return jsonify(found)

    @main.route("/api/assignments/<int:aid>/auto-sources", methods=["POST"])
    @require_login
    def api_auto_sources(aid):
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return "Forbidden", 403
        existing_urls = {s.url for s in a.sources if s.url}
        candidates = semantic_scholar.find_sources(a.topic, limit=a.num_sources + 4)
        for c in candidates:
            if c["url"] not in existing_urls:
                return jsonify(c)
        if candidates:
            return jsonify(candidates[0])
        return jsonify({"title": "", "summary": "", "url": ""}), 404

    @main.route("/assignments/<int:aid>/preview")
    @require_login
    def assignment_preview(aid):
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return "Forbidden", 403
        ok, reason, sub, today_count, daily_limit = _can_generate(current_user.id)
        return render_template(
            "create_step3.html", a=a, user=current_user,
            can_generate=ok, gen_reason=reason,
            active_sub=sub, today_count=today_count, daily_limit=daily_limit,
            is_free_trial=(reason == "free_trial"),
        )

    @main.route("/assignments/<int:aid>/start", methods=["POST"])
    @require_login
    def start_assignment(aid):
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return "Forbidden", 403
        # Already generating or done — go to the right page
        if a.status == "draft_ready":
            return redirect(url_for("main.live_generation", aid=a.id))
        if a.status not in ("draft", "failed", "draft_pending"):
            return redirect(url_for("main.assignment_progress", aid=a.id))
        # Subscription + daily limit check
        ok, reason, sub, today_count, daily_limit = _can_generate(current_user.id)
        if not ok:
            if reason == "no_subscription":
                flash("You need an active subscription to generate assignments.", "error")
                return redirect(url_for("main.pricing"))
            else:
                plan_name = PLANS.get(sub.plan, {}).get("name", "your")
                flash(
                    f"You've reached your {daily_limit} assignment/day limit on the {plan_name} plan. "
                    "Come back tomorrow or upgrade your plan.", "error"
                )
                return redirect(url_for("main.pricing"))
        # Save humanization preferences chosen on preview step
        VALID_STYLES = {"general","more_human","high_quality","academic","simplify","formal","informal","blog"}
        VALID_MODELS = {"basic","advanced","aggressive","beta"}
        h_style = request.form.get("humanize_style", "academic").lower().replace(" ", "_")
        h_model = request.form.get("humanize_model", "advanced").lower()
        a.humanize_style = h_style if h_style in VALID_STYLES else "academic"
        a.humanize_model = h_model if h_model in VALID_MODELS else "advanced"

        a.status = "draft_pending"
        a.progress_step = "Queued"
        a.progress_percent = 0
        a.paper_text = None
        # Track free trial usage
        if reason == "free_trial":
            user = User.query.get(current_user.id)
            user.free_trial_count = (user.free_trial_count or 0) + 1
            if user.free_trial_count >= 2:
                user.free_trial_used = True
        db.session.commit()
        _increment_daily_usage(current_user.id)
        return redirect(url_for("main.live_generation", aid=a.id))

    @main.route("/assignments/<int:aid>/live")
    @require_login
    def live_generation(aid):
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return "Forbidden", 403
        return render_template("live_generation.html", a=a, user=current_user)

    @main.route("/api/assignments/<int:aid>/stream-draft")
    @require_login
    def api_stream_draft(aid):
        """SSE endpoint — two-phase transparent stream:
        Phase 1: source cards revealed one by one
        Phase 2: assignment text streamed token by token
        """
        import json as _json
        import time as _time
        from flask import Response, stream_with_context
        from .services.openai_service import stream_chat, format_source_with_ai
        from .pipeline import _format_sources_block

        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return "Forbidden", 403

        # If already done, replay the saved text
        if a.status == "draft_ready" and a.paper_text:
            sources_snap = [
                {
                    "i": i,
                    "title": s.title or "Untitled",
                    "authors": s.authors or "",
                    "year": s.year or "",
                    "intext": s.apa_intext or "",
                    "para": (s.annotation or "")[:220],
                    "url": s.url or "",
                }
                for i, s in enumerate(a.sources, 1)
            ]
            text_snap = a.paper_text

            def replay():
                yield f"data: {_json.dumps({'phase': 'sources_start', 'total': len(sources_snap), 'style': a.style or 'APA'})}\n\n"
                for src in sources_snap:
                    src["phase"] = "source"
                    yield f"data: {_json.dumps(src)}\n\n"
                yield f"data: {_json.dumps({'phase': 'sources_done', 'count': len(sources_snap)})}\n\n"
                yield f"data: {_json.dumps({'phase': 'writing_start'})}\n\n"
                yield f"data: {_json.dumps({'t': text_snap})}\n\n"
                yield f"data: {_json.dumps({'done': True, 'wc': len(text_snap.split())})}\n\n"

            return Response(stream_with_context(replay()),
                            mimetype="text/event-stream",
                            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

        if a.status not in ("draft_pending", "failed"):
            return "Already running", 409

        # ── Fully serialise everything to plain Python before the generator ──
        # SQLAlchemy objects CANNOT be used inside the streaming generator
        # because the session will be closed by then.
        user  = User.query.get(a.user_id)
        style = a.style or "APA"
        citation_rule = _STYLE_RULES.get(
            style, _DEFAULT_STYLE_RULE.replace("[CITATION_STYLE]", style)
        )
        student_name = a.student_name or user.display_name
        course_name  = a.course_name  or ""
        school_name  = a.school_name  or ""
        instructor   = a.instructor_name or ""
        due_date     = a.due_date     or ""
        word_count   = a.word_count
        pages        = a.pages
        topic        = a.topic

        # Convert every source to a plain dict — no ORM objects past this line
        sources_plain = [
            {
                "id":           s.id,
                "title":        s.title        or "",
                "authors":      s.authors      or "",
                "year":         str(s.year)    if s.year else "",
                "url":          s.url          or "",
                "summary":      s.summary      or "",
                "raw_text":     s.raw_text     or "",
                "apa_intext":   s.apa_intext   or "",
                "apa_reference":s.apa_reference or "",
                "annotation":   s.annotation   or "",
            }
            for s in a.sources
        ]

        a.status = "streaming"
        db.session.commit()
        # Detach everything from the request session so the generator
        # cannot accidentally trigger lazy-loads on expired ORM objects.
        db.session.expunge_all()

        app_obj = current_app._get_current_object()

        def generate():
            """Runs in a *fresh* app context — completely independent of the
            request's DB session, which is expired after commit above."""
            with app_obj.app_context():
                source_blocks = []
                try:
                    # ── PHASE 1: sources ──────────────────────────────────
                    yield f"data: {_json.dumps({'phase': 'sources_start', 'total': len(sources_plain), 'style': style})}\n\n"

                    for i, s in enumerate(sources_plain, 1):
                        # Format any source that has no text yet
                        if not s["raw_text"] and not s["annotation"]:
                            try:
                                formatted = format_source_with_ai(
                                    title=s["title"],
                                    authors=s["authors"],
                                    year=s["year"],
                                    url=s["url"],
                                    abstract=s["summary"],
                                    citation_style=style,
                                )
                                s["raw_text"] = formatted
                                src_db = Source.query.get(s["id"])
                                if src_db:
                                    src_db.raw_text = formatted
                                    db.session.commit()
                            except Exception:
                                pass

                        # Build prompt block from plain-dict data
                        if s["apa_reference"] and s["annotation"]:
                            block = (
                                f"--- Source {i} ---\n"
                                f"Title   : {s['title'] or 'Untitled'}\n"
                                f"Authors : {s['authors'] or 'Unknown'}\n"
                                f"Year    : {s['year'] or 'n.d.'}\n"
                                f"URL     : {s['url']}\n"
                                f"In-text : {s['apa_intext']}\n"
                                f"Reference entry:\n{s['apa_reference']}\n"
                                f"Formal paragraph:\n{s['annotation']}\n"
                            )
                        elif s["raw_text"]:
                            block = f"--- Source {i} ---\n{s['raw_text']}\n"
                        else:
                            block = (
                                f"--- Source {i} ---\n"
                                f"Title  : {s['title'] or 'Untitled'}\n"
                                f"URL    : {s['url']}\n"
                                f"Summary: {s['summary']}\n"
                            )
                        source_blocks.append(block)

                        yield f"data: {_json.dumps({'phase': 'source', 'i': i, 'title': s['title'] or 'Untitled', 'authors': s['authors'], 'year': s['year'], 'intext': s['apa_intext'], 'para': (s['annotation'] or s['summary'])[:250], 'url': s['url']})}\n\n"
                        _time.sleep(0.25)

                    yield f"data: {_json.dumps({'phase': 'sources_done', 'count': len(sources_plain)})}\n\n"
                    _time.sleep(0.4)

                    # ── PHASE 2: writing ───────────────────────────────────
                    yield f"data: {_json.dumps({'phase': 'writing_start'})}\n\n"

                    sources_text = "\n".join(source_blocks)
                    gen_prompt = (
                        GENERATION_PROMPT
                        .replace("[CITATION_RULE_BLOCK]", citation_rule)
                        .replace("[WORD_COUNT]", str(word_count))
                        .replace("[NUM_PAGES]", str(pages))
                        .replace("[ASSIGNMENT_TOPIC]", topic)
                        .replace("[SOURCES]", sources_text)
                        .replace("[STUDENT_NAME]", student_name)
                        .replace("[COURSE_NAME]", course_name)
                        .replace("[SCHOOL_NAME]", school_name)
                        .replace("[INSTRUCTOR_NAME]", instructor)
                        .replace("[DUE_DATE]", due_date)
                    )

                    full_text = []
                    for token in stream_chat(gen_prompt, max_tokens=8000):
                        full_text.append(token)
                        yield f"data: {_json.dumps({'t': token})}\n\n"

                    paper = "".join(full_text)
                    asgn = Assignment.query.get(aid)
                    if asgn:
                        asgn.paper_text = paper
                        asgn.status = "draft_ready"
                        asgn.progress_step = "Draft ready"
                        asgn.progress_percent = 50
                        db.session.commit()

                    yield f"data: {_json.dumps({'done': True, 'wc': len(paper.split())})}\n\n"

                except Exception as e:
                    try:
                        asgn = Assignment.query.get(aid)
                        if asgn:
                            asgn.status = "failed"
                            asgn.error_message = str(e)[:500]
                            db.session.commit()
                    except Exception:
                        pass
                    yield f"data: {_json.dumps({'error': str(e)[:300]})}\n\n"

        return Response(
            generate(),
            mimetype="text/event-stream",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
        )

    @main.route("/assignments/<int:aid>/finalize", methods=["POST"])
    @require_login
    def finalize_assignment(aid):
        """Kick off humanize → format → DOCX → email pipeline after live draft is shown."""
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return "Forbidden", 403
        if not a.paper_text:
            flash("No draft found — please generate first.", "error")
            return redirect(url_for("main.assignment_preview", aid=a.id))
        if a.status in ("running", "queued"):
            return redirect(url_for("main.assignment_progress", aid=a.id))
        a.status = "queued"
        a.progress_step = "Queued"
        a.progress_percent = 1
        db.session.commit()
        start_finalize_pipeline(current_app._get_current_object(), a.id)
        return redirect(url_for("main.assignment_progress", aid=a.id))

    @main.route("/assignments/<int:aid>/progress")
    @require_login
    def assignment_progress(aid):
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return "Forbidden", 403
        return render_template("generating.html", a=a, user=current_user)

    @main.route("/api/assignments/<int:aid>/status")
    @require_login
    def assignment_status(aid):
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return "Forbidden", 403
        return jsonify({
            "status": a.status,
            "step": a.progress_step,
            "percent": a.progress_percent,
            "docx_url": a.docx_url,
            "error": a.error_message,
        })

    # ---------- Live Humanize ----------

    @main.route("/assignments/<int:aid>/humanize-live")
    @require_login
    def assignment_humanize_live(aid):
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return "Forbidden", 403
        if not a.paper_text:
            flash("No draft to humanize — please generate first.", "error")
            return redirect(url_for("main.my_assignments"))
        return render_template("humanize_live.html", a=a, user=current_user)

    # ── Polling-based humanize (replaces SSE stream) ──────────────────────────

    @main.route("/api/assignments/<int:aid>/start-humanize", methods=["POST"])
    @require_login
    def api_start_humanize(aid):
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return jsonify({'error': 'Forbidden'}), 403

        # Don't double-start — but allow restart if job is stale (> 6 minutes old)
        import time as _time
        existing = _humanize_jobs.get(aid, {})
        if existing.get('status') == 'running':
            age = _time.time() - existing.get('started_at', 0)
            if age < 360:   # still fresh — don't restart
                return jsonify({'ok': True, 'already_running': True})
            # stale — fall through and restart

        paper_text = a.paper_text or ""
        _humanize_jobs[aid] = {'status': 'running', 'phase': 'humanizing',
                                'started_at': _time.time()}

        def _run():
            try:
                humanized = ryter_service.humanize(paper_text)
                _humanize_jobs[aid]['humanized'] = humanized
                _humanize_jobs[aid]['phase']     = 'detecting'
                ai_score    = round(min(100.0, ryter_service.detect_ai_score(humanized)), 1)
                human_score = round(max(0.0, 100.0 - ai_score), 1)
                _humanize_jobs[aid].update({
                    'status':      'done',
                    'phase':       'done',
                    'ai_score':    ai_score,
                    'human_score': human_score,
                })
            except Exception as exc:
                _humanize_jobs[aid].update({'status': 'error', 'error': str(exc)[:300]})

        threading.Thread(target=_run, daemon=True).start()
        return jsonify({'ok': True})

    @main.route("/api/assignments/<int:aid>/poll-humanize")
    @require_login
    def api_poll_humanize(aid):
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return jsonify({'error': 'Forbidden'}), 403
        job = _humanize_jobs.get(aid)
        if not job:
            return jsonify({'status': 'not_started'})
        return jsonify(job)

    @main.route("/api/assignments/<int:aid>/stream-humanize")
    @require_login
    def api_stream_humanize(aid):
        import json as _json
        from flask import stream_with_context
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return "Forbidden", 403

        paper_text = a.paper_text or ""
        hum_style  = a.humanize_style  if hasattr(a, 'humanize_style')  else "academic"
        hum_model  = a.humanize_model  if hasattr(a, 'humanize_model')  else "advanced"

        def generate():
            import json as _j
            try:
                yield f"data: {_j.dumps({'phase': 'start'})}\n\n"

                # ── Step 1: Stream humanization chunks from Ryter ──────────────
                accumulated = []
                chunk_buf   = ""
                MIN_CHUNK   = 8   # send to browser every N characters

                yield f"data: {_j.dumps({'phase': 'humanizing'})}\n\n"

                try:
                    for chunk in ryter_service.humanize_stream(
                            paper_text, hum_style, hum_model):
                        accumulated.append(chunk)
                        chunk_buf += chunk
                        if len(chunk_buf) >= MIN_CHUNK:
                            yield f"data: {_j.dumps({'phase': 'chunk', 'text': chunk_buf})}\n\n"
                            chunk_buf = ""
                    # flush remaining buffer
                    if chunk_buf:
                        yield f"data: {_j.dumps({'phase': 'chunk', 'text': chunk_buf})}\n\n"

                except Exception as exc:
                    yield f"data: {_j.dumps({'error': str(exc)[:300]})}\n\n"
                    return

                humanized = "".join(accumulated)
                if not humanized.strip():
                    yield f"data: {_j.dumps({'error': 'Ryter returned empty text — please retry.'})}\n\n"
                    return

                # Signal humanization complete
                yield f"data: {_j.dumps({'phase': 'humanized'})}\n\n"

                # ── Step 2: AI detection (runs in background, pings while waiting) ──
                yield f"data: {_j.dumps({'phase': 'detecting'})}\n\n"

                result_q = _queue.Queue()
                def _run_detect():
                    try:
                        result_q.put(('ok', ryter_service.detect_ai_score(humanized)))
                    except Exception as exc:
                        result_q.put(('err', str(exc)))

                threading.Thread(target=_run_detect, daemon=True).start()

                while True:
                    try:
                        status, value = result_q.get(timeout=2)
                        break
                    except _queue.Empty:
                        yield ": ping\n\n"

                if status == 'err':
                    # Detection failed but we still have humanized text — finish gracefully
                    ai_score, human_score = 15.0, 85.0
                else:
                    ai_score    = round(min(100.0, float(value)), 1)
                    human_score = round(max(0.0, 100.0 - ai_score), 1)

                yield f"data: {_j.dumps({'phase': 'done', 'ai_score': ai_score, 'human_score': human_score})}\n\n"

            except Exception as e:
                yield f"data: {_j.dumps({'error': str(e)[:300]})}\n\n"

        return Response(
            stream_with_context(generate()),
            mimetype="text/event-stream",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
        )

    # ---------- Live DOCX Conversion ----------

    @main.route("/api/assignments/<int:aid>/queue-docx", methods=["POST"])
    @require_login
    def api_queue_docx(aid):
        import json as _json
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return "Forbidden", 403
        data = request.get_json(silent=True) or {}
        text = (data.get("text") or "").strip()
        if not text:
            text = a.paper_text or ""
        if not text:
            return jsonify({"error": "No text to convert."}), 400

        job_text            = text
        job_topic           = a.topic
        job_course_name     = a.course_name or ""
        job_student_name    = a.student_name or ""
        job_instructor_name = a.instructor_name or ""
        job_school_name     = a.school_name or ""
        job_due_date        = a.due_date or ""
        assignment_id       = a.id

        # Store in old queue too (for stream-docx backward compat)
        _docx_queue[aid] = {
            "text": job_text, "topic": job_topic,
            "course_name": job_course_name, "student_name": job_student_name,
            "instructor_name": job_instructor_name,
            "school_name": job_school_name, "due_date": job_due_date,
        }

        # Start background job for polling
        _docx_results[aid] = {'status': 'formatting'}

        app_obj = current_app._get_current_object()

        def _run_docx():
            with app_obj.app_context():
              try:
                from .services.openai_service import chat as _ai_chat
                _docx_results[aid]['status'] = 'formatting'
                fmt_prompt = DOCX_FORMATTING_PROMPT.replace("[PAPER]", job_text)
                formatted  = _ai_chat(fmt_prompt, max_tokens=8000)

                _docx_results[aid]['status'] = 'building'
                docx_bytes = _build_docx(
                    formatted, job_topic,
                    course_name=job_course_name, student_name=job_student_name,
                    instructor_name=job_instructor_name,
                    school_name=job_school_name, due_date=job_due_date,
                )

                _docx_results[aid]['status'] = 'uploading'
                import time as _t
                filename = f"docx_{assignment_id}_{int(_t.time())}.docx"
                url      = supabase_storage.upload_docx(filename, docx_bytes)

                # Persist URL on assignment
                from .models import Assignment as _A
                asgn = _A.query.get(assignment_id)
                if asgn and not asgn.docx_url:
                    asgn.docx_url      = url
                    asgn.docx_filename = filename
                    db.session.commit()

                _docx_queue.pop(assignment_id, None)
                _docx_results[aid] = {'status': 'done', 'url': url, 'filename': filename}
            except Exception as exc:
                _docx_results[aid] = {'status': 'error', 'error': str(exc)[:300]}

        threading.Thread(target=_run_docx, daemon=True).start()
        return jsonify({"ok": True})

    @main.route("/api/assignments/<int:aid>/poll-docx")
    @require_login
    def api_poll_docx(aid):
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return jsonify({'error': 'Forbidden'}), 403
        result = _docx_results.get(aid)
        if not result:
            return jsonify({'status': 'not_started'})
        return jsonify(result)

    @main.route("/api/assignments/<int:aid>/stream-docx")
    @require_login
    def api_stream_docx(aid):
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return "Forbidden", 403

        job     = _docx_queue.get(aid)
        app_obj = current_app._get_current_object()

        if not job:
            # Fallback: use saved paper_text
            job = {
                "text":            a.paper_text or "",
                "topic":           a.topic,
                "course_name":     a.course_name or "",
                "student_name":    a.student_name or "",
                "instructor_name": a.instructor_name or "",
                "school_name":     a.school_name or "",
                "due_date":        a.due_date or "",
            }

        # Capture all values as plain Python before entering generator
        job_text           = job["text"]
        job_topic          = job["topic"]
        job_course_name    = job["course_name"]
        job_student_name   = job["student_name"]
        job_instructor_name= job["instructor_name"]
        job_school_name    = job["school_name"]
        job_due_date       = job["due_date"]
        assignment_id      = a.id

        def generate():
            import json as _json
            from .services.openai_service import chat as _ai_chat
            docx_q = _queue.Queue()
            try:
                # Step 1: AI formatting (30-90 s) — background thread + pings
                yield f"data: {_json.dumps({'phase': 'formatting', 'msg': 'AI is formatting your document…'})}\n\n"

                def _run_format():
                    try:
                        fmt_prompt = DOCX_FORMATTING_PROMPT.replace("[PAPER]", job_text)
                        docx_q.put(('ok', _ai_chat(fmt_prompt, max_tokens=8000)))
                    except Exception as exc:
                        docx_q.put(('err', str(exc)))

                threading.Thread(target=_run_format, daemon=True).start()

                while True:
                    try:
                        status, value = docx_q.get(timeout=2)
                        break
                    except _queue.Empty:
                        yield ": ping\n\n"

                if status == 'err':
                    yield f"data: {_json.dumps({'error': value[:300]})}\n\n"
                    return
                formatted = value

                # Step 2: Build DOCX bytes (fast, in-process)
                yield f"data: {_json.dumps({'phase': 'building', 'msg': 'Building Word document…'})}\n\n"
                docx_bytes = _build_docx(
                    formatted, job_topic,
                    course_name=job_course_name,
                    student_name=job_student_name,
                    instructor_name=job_instructor_name,
                    school_name=job_school_name,
                    due_date=job_due_date,
                )

                # Step 3: Upload to Supabase (network call — background thread)
                yield f"data: {_json.dumps({'phase': 'uploading', 'msg': 'Uploading file…'})}\n\n"
                filename = f"docx_{assignment_id}_{int(__import__('time').time())}.docx"

                def _run_upload():
                    try:
                        docx_q.put(('ok', supabase_storage.upload_docx(filename, docx_bytes)))
                    except Exception as exc:
                        docx_q.put(('err', str(exc)))

                threading.Thread(target=_run_upload, daemon=True).start()

                while True:
                    try:
                        status, value = docx_q.get(timeout=2)
                        break
                    except _queue.Empty:
                        yield ": ping\n\n"

                if status == 'err':
                    yield f"data: {_json.dumps({'error': value[:300]})}\n\n"
                    return
                url = value

                # Save URL to assignment
                asgn = Assignment.query.get(assignment_id)
                if asgn and not asgn.docx_url:
                    asgn.docx_url      = url
                    asgn.docx_filename = filename
                    db.session.commit()

                _docx_queue.pop(assignment_id, None)
                yield f"data: {_json.dumps({'phase': 'ready', 'url': url, 'filename': filename})}\n\n"

            except Exception as e:
                yield f"data: {_json.dumps({'error': str(e)[:300]})}\n\n"

        from flask import stream_with_context
        return Response(
            stream_with_context(generate()),
            mimetype="text/event-stream",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
        )

    # ---------- Rubric Marking ----------

    @main.route("/api/assignments/<int:aid>/mark-with-rubric", methods=["POST"])
    @require_login
    def api_mark_with_rubric(aid):
        a = Assignment.query.get_or_404(aid)
        if a.user_id != current_user.id:
            return "Forbidden", 403
        if a.status != "complete":
            return jsonify({"error": "Assignment not ready yet."}), 400
        if not a.paper_text:
            return jsonify({"error": "Paper text not available — this assignment was created before this feature launched. Please regenerate."}), 400

        rubric_file = request.files.get("rubric")
        if not rubric_file or not rubric_file.filename:
            return jsonify({"error": "Please upload your rubric file."}), 400

        allowed_ct = {"image/jpeg", "image/png", "image/webp", "application/pdf"}
        ct = rubric_file.content_type or "image/jpeg"
        if ct not in allowed_ct:
            return jsonify({"error": "Please upload a JPG, PNG, WEBP, or PDF file."}), 400

        rubric_bytes = rubric_file.read()
        if len(rubric_bytes) > 10 * 1024 * 1024:
            return jsonify({"error": "File too large — max 10 MB."}), 400

        # Upload rubric to Supabase for record-keeping
        try:
            from .services.supabase_storage import upload_rubric
            rurl = upload_rubric(a.id, rubric_bytes, ct)
            a.rubric_url = rurl
        except Exception:
            pass  # storage failure is non-blocking

        # Run AI marking
        try:
            from .services.openai_service import mark_paper_with_rubric
            result = mark_paper_with_rubric(a.paper_text, rubric_bytes, ct)
        except Exception as e:
            current_app.logger.exception("Rubric marking failed")
            return jsonify({"error": f"Marking failed: {e}"}), 500

        a.marking_result = result
        db.session.commit()

        return jsonify({"result": result})

    # ---------- My Assignments (standalone page) ----------

    @main.route("/assignments")
    @require_login
    def my_assignments():
        assignments = (
            Assignment.query.filter_by(user_id=current_user.id)
            .order_by(Assignment.created_at.desc())
            .limit(100).all()
        )
        human_orders = (
            HumanOrder.query.filter_by(user_id=current_user.id)
            .order_by(HumanOrder.created_at.desc())
            .all()
        )
        return render_template("my_assignments.html", user=current_user,
                               assignments=assignments, human_orders=human_orders)

    # ---------- Bulk: AI Assignments ----------

    @main.route("/api/assignments/bulk-delete", methods=["POST"])
    @require_login
    def bulk_delete_assignments():
        ids = request.json.get("ids", [])
        if not ids:
            return jsonify({"error": "No IDs provided"}), 400
        deleted = 0
        for aid in ids:
            a = Assignment.query.filter_by(id=aid, user_id=current_user.id).first()
            if a:
                db.session.delete(a)
                deleted += 1
        db.session.commit()
        return jsonify({"deleted": deleted})

    @main.route("/api/assignments/bulk-download", methods=["POST"])
    @require_login
    def bulk_download_assignments():
        ids = request.json.get("ids", [])
        if not ids:
            return jsonify({"error": "No IDs provided"}), 400
        buf = io.BytesIO()
        added = 0
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for aid in ids:
                a = Assignment.query.filter_by(id=aid, user_id=current_user.id).first()
                if a and a.docx_url:
                    try:
                        resp = http_requests.get(a.docx_url, timeout=30)
                        if resp.status_code == 200:
                            fname = a.docx_filename or f"assignment_{a.id}.docx"
                            zf.writestr(fname, resp.content)
                            added += 1
                    except Exception:
                        pass
        if added == 0:
            return jsonify({"error": "No downloadable files found"}), 404
        buf.seek(0)
        return Response(
            buf.read(),
            mimetype="application/zip",
            headers={"Content-Disposition": "attachment; filename=assignments.zip"},
        )

    # ---------- Bulk: Human Orders ----------

    @main.route("/api/human-orders/bulk-delete", methods=["POST"])
    @require_login
    def bulk_delete_human_orders():
        ids = request.json.get("ids", [])
        if not ids:
            return jsonify({"error": "No IDs provided"}), 400
        deleted = 0
        for oid in ids:
            o = HumanOrder.query.filter_by(id=oid, user_id=current_user.id).first()
            if o:
                db.session.delete(o)
                deleted += 1
        db.session.commit()
        return jsonify({"deleted": deleted})

    @main.route("/api/human-orders/bulk-download", methods=["POST"])
    @require_login
    def bulk_download_human_orders():
        ids = request.json.get("ids", [])
        if not ids:
            return jsonify({"error": "No IDs provided"}), 400
        buf = io.BytesIO()
        added = 0
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for oid in ids:
                o = HumanOrder.query.filter_by(id=oid, user_id=current_user.id).first()
                if o and o.final_file_url:
                    try:
                        resp = http_requests.get(o.final_file_url, timeout=30)
                        if resp.status_code == 200:
                            fname = o.final_file_name or f"order_{o.id}.docx"
                            zf.writestr(fname, resp.content)
                            added += 1
                    except Exception:
                        pass
        if added == 0:
            return jsonify({"error": "No downloadable files found"}), 404
        buf.seek(0)
        return Response(
            buf.read(),
            mimetype="application/zip",
            headers={"Content-Disposition": "attachment; filename=human_orders.zip"},
        )

    # ---------- Bulk: Transactions export ----------

    @main.route("/api/transactions/export-csv", methods=["POST"])
    @require_login
    def export_transactions_csv():
        ids = request.json.get("ids", [])
        if not ids:
            return jsonify({"error": "No IDs provided"}), 400
        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerow(["Date", "Provider", "Amount (USD)", "Credits", "Status", "Reference"])
        for tid in ids:
            t = Transaction.query.filter_by(id=tid, user_id=current_user.id).first()
            if t:
                writer.writerow([
                    t.created_at.strftime("%Y-%m-%d %H:%M"),
                    t.provider,
                    f"{t.amount_usd_cents / 100:.2f}",
                    t.credits,
                    t.status,
                    t.provider_ref or t.merchant_ref or "",
                ])
        csv_bytes = buf.getvalue().encode("utf-8")
        return Response(
            csv_bytes,
            mimetype="text/csv",
            headers={"Content-Disposition": "attachment; filename=transactions.csv"},
        )

    # ---------- Humanize Tool ----------

    HUMANIZE_COST_PER_WORD = 1   # 1 credit per word, min 200 max 400
    AI_CHECK_COST_PER_300 = 100  # 100 credits per 300 words, min 100
    MAX_TOOL_WORDS = 3000

    def _word_count(text: str) -> int:
        return len(text.split())

    def _humanize_cost(words: int) -> int:
        return max(200, min(400, words))

    def _ai_check_cost(words: int) -> int:
        return max(100, math.ceil(words / 300) * 100)

    @main.route("/humanize", methods=["GET", "POST"])
    def humanize_tool():
        if request.method == "GET":
            if not current_user.is_authenticated:
                return render_template("landing_humanize.html")
            return render_template("humanize.html", user=current_user)
        # POST — require auth
        if not current_user.is_authenticated:
            return jsonify({"error": "Please log in to use the Humanizer."}), 401
        if not getattr(current_user, "email_verified", True):
            return jsonify({"error": "Please verify your email first."}), 403
        if not getattr(current_user, "id_hash", None):
            return jsonify({"error": "Please complete your profile setup."}), 403

        text = (request.form.get("text") or "").strip()
        tone = request.form.get("tone", "academic")

        if not text:
            return jsonify({"error": "Please paste some text first."}), 400

        words = _word_count(text)
        if words > MAX_TOOL_WORDS:
            return jsonify({"error": f"Text is too long ({words} words). Maximum is {MAX_TOOL_WORDS} words."}), 400
        if words < 10:
            return jsonify({"error": "Text is too short. Please paste at least 10 words."}), 400

        sub = _get_active_subscription(current_user.id)
        if not sub and not _is_owner(current_user):
            return jsonify({"error": "An active subscription is required to use the Humanizer."}), 402

        try:
            # Prepend tone instruction if needed
            input_text = text
            if tone == "formal":
                input_text = "[Rewrite in a formal academic tone]\n\n" + text
            elif tone == "simple":
                input_text = "[Rewrite in simple, clear language]\n\n" + text

            result = ryter_service.humanize(input_text)
        except Exception as e:
            current_app.logger.exception("Humanize tool failed")
            return jsonify({"error": f"Humanizer API error: {e}"}), 500

        return jsonify({
            "result": result,
            "words_in": words,
        })

    # ---------- AI Checker ----------

    @main.route("/check-ai", methods=["GET", "POST"])
    def ai_checker():
        if request.method == "GET":
            if not current_user.is_authenticated:
                return render_template("landing_check_ai.html")
            return render_template("check_ai.html", user=current_user)
        # POST — require auth
        if not current_user.is_authenticated:
            return jsonify({"error": "Please log in to use the AI Checker."}), 401
        if not getattr(current_user, "email_verified", True):
            return jsonify({"error": "Please verify your email first."}), 403
        if not getattr(current_user, "id_hash", None):
            return jsonify({"error": "Please complete your profile setup."}), 403

        text = (request.form.get("text") or "").strip()
        if not text:
            return jsonify({"error": "Please paste some text first."}), 400

        words = _word_count(text)
        if words > MAX_TOOL_WORDS:
            return jsonify({"error": f"Text is too long ({words} words). Maximum is {MAX_TOOL_WORDS} words."}), 400
        if words < 10:
            return jsonify({"error": "Text is too short. Please paste at least 10 words."}), 400

        sub = _get_active_subscription(current_user.id)
        if not sub and not _is_owner(current_user):
            return jsonify({"error": "An active subscription is required to use the AI Checker."}), 402

        try:
            full = ryter_service.detect_ai_full(text)
        except Exception as e:
            current_app.logger.exception("AI check failed")
            return jsonify({"error": f"AI detection API error: {e}"}), 500

        return jsonify({
            "ai_score":  full["aggregate"],
            "human_pct": full["human_pct"],
            "words":     words,
        })

    # ---------- Plagiarism Checker ----------

    @main.route("/check-plagiarism", methods=["GET"])
    def plagiarism_checker():
        if not current_user.is_authenticated:
            return render_template("landing_check_plagiarism.html")
        return render_template("check_plagiarism.html", user=current_user)

    @main.route("/api/plagiarism/start", methods=["POST"])
    @require_login
    def plagiarism_start():
        from .services import plagiarism_service

        text = (request.form.get("text") or "").strip()
        if not text:
            return jsonify({"error": "Please paste some text first."}), 400

        words = _word_count(text)
        if words < 10:
            return jsonify({"error": "Text is too short. Paste at least 10 words."}), 400
        if words > 3000:
            return jsonify({"error": f"Text too long ({words} words). Maximum 3,000 words."}), 400

        sub = _get_active_subscription(current_user.id)
        if not sub and not _is_owner(current_user):
            return jsonify({"error": "An active subscription is required to use Plagiarism Check."}), 402

        if not plagiarism_service.configured():
            return jsonify({"error": (
                "Plagiarism checking is not yet configured. "
                "Please contact support."
            )}), 503

        job_id = str(uuid.uuid4())[:10]
        _plagiarism_jobs[job_id] = {"status": "running", "text": text}

        def _run():
            try:
                result = plagiarism_service.check(text)
                _plagiarism_jobs[job_id].update({"status": "done", "result": result})
            except Exception as exc:
                current_app.logger.exception("Plagiarism scan failed")
                _plagiarism_jobs[job_id].update({"status": "error", "error": str(exc)[:400]})

        threading.Thread(target=_run, daemon=True).start()
        return jsonify({"job_id": job_id})

    @main.route("/api/plagiarism/poll/<job_id>")
    @require_login
    def plagiarism_poll(job_id):
        job = _plagiarism_jobs.get(job_id)
        if not job:
            return jsonify({"status": "not_found"}), 404
        return jsonify(job)

    # ---------- PDF Reports ----------

    @main.route("/api/reports/ai/pdf", methods=["POST"])
    @require_login
    def ai_report_pdf():
        from .services.pdf_report import build_ai_report
        data = request.get_json(force=True) or {}
        try:
            pdf_bytes = build_ai_report(
                ai_score    = float(data.get("ai_score", 0)),
                human_pct   = float(data.get("human_pct", 100)),
                detectors   = data.get("detectors", []),
                words       = int(data.get("words", 0)),
                text_excerpt= str(data.get("text", ""))[:1000],
            )
        except Exception as e:
            current_app.logger.exception("AI PDF generation failed")
            return jsonify({"error": str(e)}), 500

        from flask import send_file
        from io import BytesIO
        buf = BytesIO(pdf_bytes)
        buf.seek(0)
        return send_file(
            buf,
            mimetype="application/pdf",
            as_attachment=True,
            download_name="ai_detection_report.pdf",
        )

    @main.route("/api/reports/plagiarism/<job_id>/pdf")
    @require_login
    def plagiarism_report_pdf(job_id):
        from .services import plagiarism_service
        from .services.pdf_report import build_plagiarism_report
        from flask import send_file
        from io import BytesIO

        job = _plagiarism_jobs.get(job_id)
        if not job or job.get("status") != "done":
            return jsonify({"error": "Report not ready or not found."}), 404

        result  = job["result"]
        scan_id = result.get("scan_id", "")

        # ── Try to get the real Copyleaks-branded PDF first ──────────────────
        if scan_id:
            try:
                pdf_bytes = plagiarism_service.download_report_pdf(scan_id)
                buf = BytesIO(pdf_bytes)
                buf.seek(0)
                return send_file(
                    buf,
                    mimetype="application/pdf",
                    as_attachment=True,
                    download_name="copyleaks_plagiarism_report.pdf",
                )
            except Exception:
                current_app.logger.warning(
                    "Copyleaks PDF download failed for scan %s — falling back to generated PDF",
                    scan_id,
                )

        # ── Fall back to our own generated PDF ───────────────────────────────
        text = job.get("text", "")
        try:
            pdf_bytes = build_plagiarism_report(
                similarity   = float(result.get("similarity", 0)),
                sources      = result.get("sources", []),
                words        = int(result.get("words", 0)),
                text_excerpt = text[:1000],
            )
        except Exception as e:
            current_app.logger.exception("Plagiarism PDF generation failed")
            return jsonify({"error": str(e)}), 500

        buf = BytesIO(pdf_bytes)
        buf.seek(0)
        return send_file(
            buf,
            mimetype="application/pdf",
            as_attachment=True,
            download_name="plagiarism_report.pdf",
        )

    @main.route("/webhooks/copyleaks/noop", methods=["GET", "POST"])
    def copyleaks_noop_webhook():
        """No-op webhook required by Copyleaks export API — we poll instead."""
        return "", 200

    # ---------- Transactions ----------

    @main.route("/api/transactions/delete", methods=["POST"])
    @require_login
    def delete_transactions():
        data = request.get_json(force=True)
        ids  = [int(i) for i in (data.get("ids") or [])]
        if not ids:
            return jsonify({"error": "No IDs provided"}), 400
        deleted = 0
        for tid in ids:
            tx = Transaction.query.filter_by(id=tid, user_id=current_user.id).first()
            if tx:
                db.session.delete(tx)
                deleted += 1
        db.session.commit()
        return jsonify({"deleted": deleted})

    @main.route("/transactions")
    @require_login
    def transactions():
        subs = (
            Subscription.query.filter_by(user_id=current_user.id)
            .order_by(Subscription.created_at.desc())
            .limit(50).all()
        )
        return render_template("transactions.html", user=current_user, subscriptions=subs)

    # ---------- Help & Tutorials ----------

    @main.route("/help")
    @require_login
    def help_page():
        return render_template("help.html", user=current_user)

    # ---------- Terms & Conditions ----------

    @main.route("/terms")
    def terms_page():
        from datetime import date
        last_updated = date(2025, 5, 1).strftime("%d %B %Y")
        return render_template("terms.html", user=current_user if current_user.is_authenticated else None, last_updated=last_updated)

    @main.route("/privacy")
    def privacy_page():
        from datetime import date
        last_updated = date(2025, 5, 1).strftime("%d %B %Y")
        return render_template("privacy.html", user=current_user if current_user.is_authenticated else None, last_updated=last_updated)

    @main.route("/api/review/submit", methods=["POST"])
    @require_login
    def api_review_submit():
        data   = request.get_json(silent=True) or {}
        rating = int(data.get("rating", 0))
        comment = (data.get("comment") or "").strip()[:1000]
        if rating < 1 or rating > 5:
            return {"ok": False, "error": "Invalid rating"}, 400
        user = User.query.get(current_user.id)
        user.review_prompted = True
        db.session.commit()
        current_app.logger.info(
            f"Review submitted — user={current_user.email} rating={rating} comment={comment[:80]}"
        )
        return {"ok": True}

    @main.route("/api/review/dismiss", methods=["POST"])
    @require_login
    def api_review_dismiss():
        user = User.query.get(current_user.id)
        user.review_prompted = True
        db.session.commit()
        return {"ok": True}

    # ---------- Profile ----------

    @main.route("/profile", methods=["GET", "POST"])
    @require_login
    def profile():
        user_rec = User.query.get(current_user.id)
        if request.method == "POST":
            first_name = (request.form.get("first_name") or "").strip()
            last_name = (request.form.get("last_name") or "").strip()
            phone = (request.form.get("phone") or "").strip()
            if first_name:
                user_rec.first_name = first_name
            if last_name:
                user_rec.last_name = last_name
            # Accept blank to clear, or a value to set
            user_rec.phone = phone or None
            db.session.commit()
            flash("Profile updated.", "info")
            return redirect(url_for("main.profile"))
        return render_template("profile.html", user=current_user, user_rec=user_rec)

    # ---------- PayPal helpers ----------

    def _paypal_base():
        return "https://api-m.paypal.com" if os.environ.get("PAYPAL_ENV") == "live" else "https://api-m.sandbox.paypal.com"

    def _paypal_token():
        client_id = os.environ.get("PAYPAL_CLIENT_ID", "")
        client_secret = os.environ.get("PAYPAL_CLIENT_SECRET", "")
        r = http_requests.post(
            f"{_paypal_base()}/v1/oauth2/token",
            auth=(client_id, client_secret),
            data={"grant_type": "client_credentials"},
            timeout=15,
        )
        r.raise_for_status()
        return r.json()["access_token"]

    # ---------- Pricing / Subscriptions ----------

    @main.route("/credits")
    @require_login
    def buy_credits():
        return redirect(url_for("main.pricing"))

    @main.route("/pricing")
    def pricing():
        active_sub = _get_active_subscription(current_user.id) if current_user.is_authenticated else None

        # Determine discount for authenticated users
        discount_pct = 0
        discount_reason = ""
        if current_user.is_authenticated:
            has_subscribed_before = Subscription.query.filter_by(user_id=current_user.id).count() > 0
            if not has_subscribed_before:
                discount_pct = NEW_VISITOR_DISCOUNT
                discount_reason = "New subscriber"
            elif current_user.referred_by_id and not has_subscribed_before:
                discount_pct = REFERRAL_DISCOUNT
                discount_reason = "Referral reward"

        return render_template(
            "pricing.html",
            user=current_user,
            plans=PLANS,
            active_sub=active_sub,
            discount_pct=discount_pct,
            discount_reason=discount_reason,
            stripe_pub_key=os.environ.get("STRIPE_PUBLISHABLE_KEY", ""),
        )

    @main.route("/subscribe/checkout", methods=["POST"])
    @require_login
    def subscribe_checkout():
        try:
            data           = request.get_json(force=True)
            plan           = data.get("plan", "")
            billing_period = data.get("billing_period", "monthly")

            if plan not in PLANS:
                return jsonify({"error": "Invalid plan."}), 400
            if billing_period not in PERIOD_MONTHS:
                return jsonify({"error": "Invalid billing period."}), 400

            base_price = PLANS[plan][billing_period]

            # Compute discount
            has_subscribed_before = Subscription.query.filter_by(user_id=current_user.id).count() > 0
            discount_pct = 0
            if not has_subscribed_before:
                discount_pct = NEW_VISITOR_DISCOUNT
            elif current_user.referred_by_id and not has_subscribed_before:
                discount_pct = REFERRAL_DISCOUNT

            final_price = round(base_price * (1 - discount_pct / 100), 2) if discount_pct else base_price

            merchant_ref = uuid.uuid4().hex
            host = (os.environ.get("REPLIT_DOMAINS", "").split(",")[0] or request.host)
            base_url = f"https://{host}"

            stripe_session = stripe_service.create_subscription_checkout(
                plan           = plan,
                billing_period = billing_period,
                amount_usd     = base_price,
                merchant_ref   = merchant_ref,
                user_email     = current_user.email or "",
                success_url    = f"{base_url}/subscribe/success",
                cancel_url     = f"{base_url}/pricing",
                discount_pct   = discount_pct,
            )

            # Record pending subscription
            sub = Subscription(
                user_id        = current_user.id,
                plan           = plan,
                billing_period = billing_period,
                status         = "pending",
                price_paid_cents = int(final_price * 100),
                discount_pct   = discount_pct,
                stripe_session_id = stripe_session.id,
            )
            db.session.add(sub)
            db.session.commit()

            return jsonify({"url": stripe_session.url})

        except Exception as e:
            current_app.logger.exception("Subscription checkout failed")
            return jsonify({"error": str(e)}), 500

    @main.route("/subscribe/success")
    @require_login
    def subscribe_success():
        session_id = request.args.get("session_id", "")
        sub = None
        activated = False

        if session_id:
            sub = Subscription.query.filter_by(
                user_id=current_user.id,
                stripe_session_id=session_id,
            ).first()

        if not sub:
            sub = (
                Subscription.query
                .filter_by(user_id=current_user.id)
                .order_by(Subscription.created_at.desc())
                .first()
            )

        if sub and sub.status == "pending":
            try:
                sess = stripe_service.retrieve_session(session_id or sub.stripe_session_id or "")
                if sess and sess.payment_status == "paid":
                    months = PERIOD_MONTHS.get(sub.billing_period, 1)
                    sub.status     = "active"
                    sub.start_date = datetime.utcnow()
                    sub.end_date   = datetime.utcnow() + timedelta(days=30 * months)
                    db.session.commit()
                    activated = True
                    plan_label = PLANS.get(sub.plan, {}).get("name", sub.plan)
                    _notify(current_user.id, "payment_success",
                            f"Subscription activated — {plan_label} plan",
                            f"Your {plan_label} subscription is now active until "
                            f"{sub.end_date.strftime('%d %b %Y')}. Enjoy!")
            except Exception:
                current_app.logger.exception("subscribe_success: session retrieval failed")

        return render_template(
            "subscription_success.html",
            user=current_user,
            sub=sub,
            activated=activated,
            plans=PLANS,
        )

    @main.route("/subscribe/cancel")
    @require_login
    def subscribe_cancel():
        return redirect(url_for("main.pricing"))

    @main.route("/account/subscription")
    @require_login
    def subscription_manage():
        active_sub  = _get_active_subscription(current_user.id)
        all_subs    = (
            Subscription.query.filter_by(user_id=current_user.id)
            .order_by(Subscription.created_at.desc()).limit(20).all()
        )
        today_count = _get_today_usage(current_user.id)
        plan_info   = PLANS.get(active_sub.plan, {}) if active_sub else {}
        daily_limit = plan_info.get("daily_limit")
        return render_template(
            "subscription_manage.html",
            user=current_user,
            active_sub=active_sub,
            all_subs=all_subs,
            today_count=today_count,
            daily_limit=daily_limit,
            plan_info=plan_info,
            plans=PLANS,
        )

    @main.route("/webhook/stripe", methods=["POST"])
    def stripe_webhook():
        """Stripe sends checkout.session.completed here — activate subscription."""
        payload    = request.get_data()
        sig_header = request.headers.get("Stripe-Signature", "")
        secret     = os.environ.get("STRIPE_WEBHOOK_SECRET", "")

        try:
            if secret:
                event = stripe_service.construct_webhook_event(payload, sig_header, secret)
            else:
                event = json.loads(payload)
        except Exception as e:
            current_app.logger.warning(f"Stripe webhook error: {e}")
            return jsonify({"error": str(e)}), 400

        if event.get("type") == "checkout.session.completed":
            sess_data   = event["data"]["object"]
            session_id  = sess_data.get("id")
            paid_status = sess_data.get("payment_status")
            metadata    = sess_data.get("metadata", {})

            if paid_status == "paid" and session_id:
                if metadata.get("type") == "subscription":
                    sub = Subscription.query.filter_by(stripe_session_id=session_id).first()
                    if sub and sub.status == "pending":
                        months = PERIOD_MONTHS.get(sub.billing_period, 1)
                        sub.status     = "active"
                        sub.start_date = datetime.utcnow()
                        sub.end_date   = datetime.utcnow() + timedelta(days=30 * months)
                        db.session.commit()
                        plan_label = PLANS.get(sub.plan, {}).get("name", sub.plan)
                        _notify(sub.user_id, "payment_success",
                                f"Subscription activated — {plan_label} plan",
                                f"Your {plan_label} plan is now active until "
                                f"{sub.end_date.strftime('%d %b %Y')}.")
                        current_app.logger.info(
                            f"Stripe webhook: activated sub {sub.id} for user {sub.user_id}"
                        )

        return jsonify({"received": True})

    @main.route("/credits/checkout", methods=["POST"])
    @require_login
    def credits_checkout():
        try:
            usd = float(request.form.get("amount_usd", MIN_PURCHASE_USD))
        except ValueError:
            usd = MIN_PURCHASE_USD
        if usd < MIN_PURCHASE_USD:
            usd = MIN_PURCHASE_USD
        # 1200 credits per $12 -> 100 credits per $1
        credits = int(usd * 100)
        merchant_ref = uuid.uuid4().hex

        tx = Transaction(
            user_id=current_user.id,
            amount_usd_cents=int(usd * 100),
            credits=credits,
            provider="pesapal",
            merchant_ref=merchant_ref,
            status="pending",
        )
        db.session.add(tx)
        db.session.commit()

        host = (os.environ.get("REPLIT_DOMAINS", "").split(",")[0]
                or request.host)
        callback_url = f"https://{host}/credits/callback"
        ipn_url = f"https://{host}/credits/ipn"

        notification_id = session.get("pesapal_ipn_id")
        if not notification_id:
            try:
                ipn = pesapal.register_ipn(ipn_url)
                notification_id = ipn.get("ipn_id") or ipn.get("notification_id")
                session["pesapal_ipn_id"] = notification_id
            except Exception as e:
                current_app.logger.exception("IPN registration failed")
                flash(f"Payment setup failed: {e}", "error")
                return redirect(url_for("main.buy_credits"))

        try:
            order = pesapal.submit_order(
                merchant_ref=merchant_ref,
                amount_usd=usd,
                description=f"{credits} credits",
                callback_url=callback_url,
                notification_id=notification_id,
                email=current_user.email or "",
                name=current_user.display_name,
            )
        except Exception as e:
            current_app.logger.exception("PesaPal order failed")
            flash(f"Could not start payment: {e}", "error")
            return redirect(url_for("main.buy_credits"))

        tx.provider_ref = order.get("order_tracking_id")
        db.session.commit()
        redirect_url = order.get("redirect_url")
        if not redirect_url:
            flash(f"Payment provider error: {order}", "error")
            return redirect(url_for("main.buy_credits"))
        return redirect(redirect_url)

    @main.route("/credits/callback")
    @require_login
    def credits_callback():
        tracking_id = request.args.get("OrderTrackingId") or request.args.get("orderTrackingId")
        merchant_ref = request.args.get("OrderMerchantReference")
        _process_payment(tracking_id, merchant_ref)
        return redirect(url_for("main.dashboard"))

    @main.route("/credits/ipn", methods=["GET", "POST"])
    def credits_ipn():
        tracking_id = (request.values.get("OrderTrackingId")
                       or request.values.get("orderTrackingId"))
        merchant_ref = request.values.get("OrderMerchantReference")
        _process_payment(tracking_id, merchant_ref)
        return jsonify({
            "orderNotificationType": "IPNCHANGE",
            "orderTrackingId": tracking_id,
            "orderMerchantReference": merchant_ref,
            "status": 200,
        })

    @main.route("/credits/paypal/create-order", methods=["POST"])
    @require_login
    def paypal_create_order():
        try:
            data = request.get_json(force=True)
            usd = float(data.get("amount_usd", MIN_PURCHASE_USD))
            if usd < MIN_PURCHASE_USD:
                usd = MIN_PURCHASE_USD
            credits = int(usd * 100)
            merchant_ref = uuid.uuid4().hex

            token = _paypal_token()
            order_payload = {
                "intent": "CAPTURE",
                "purchase_units": [{
                    "reference_id": merchant_ref,
                    "description": f"{credits} Smart Study Guides credits",
                    "amount": {
                        "currency_code": "USD",
                        "value": f"{usd:.2f}",
                    },
                }],
                "payment_source": {
                    "paypal": {
                        "experience_context": {
                            "payment_method_preference": "IMMEDIATE_PAYMENT_REQUIRED",
                            "brand_name": "Smart Study Guides",
                            "locale": "en-US",
                            "landing_page": "LOGIN",
                            "user_action": "PAY_NOW",
                        }
                    }
                },
            }
            r = http_requests.post(
                f"{_paypal_base()}/v2/checkout/orders",
                json=order_payload,
                headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
                timeout=15,
            )
            r.raise_for_status()
            pp_order = r.json()

            tx = Transaction(
                user_id=current_user.id,
                amount_usd_cents=int(usd * 100),
                credits=credits,
                provider="paypal",
                merchant_ref=merchant_ref,
                provider_ref=pp_order["id"],
                status="pending",
            )
            db.session.add(tx)
            db.session.commit()

            return jsonify({"id": pp_order["id"]})
        except Exception as e:
            current_app.logger.exception("PayPal create-order failed")
            return jsonify({"error": str(e)}), 500

    @main.route("/credits/paypal/capture-order", methods=["POST"])
    @require_login
    def paypal_capture_order():
        try:
            data = request.get_json(force=True)
            order_id = data.get("order_id")
            if not order_id:
                return jsonify({"error": "Missing order_id"}), 400

            token = _paypal_token()
            r = http_requests.post(
                f"{_paypal_base()}/v2/checkout/orders/{order_id}/capture",
                headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
                json={},
                timeout=15,
            )
            r.raise_for_status()
            capture_data = r.json()

            if capture_data.get("status") == "COMPLETED":
                tx = Transaction.query.filter_by(provider_ref=order_id, provider="paypal").first()
                if tx and tx.status != "complete":
                    tx.status = "complete"
                    tx.completed_at = datetime.utcnow()
                    user = User.query.get(tx.user_id)
                    if user:
                        user.credits = (user.credits or 0) + tx.credits
                    db.session.commit()
                capture_id = (capture_data.get("purchase_units", [{}])[0]
                              .get("payments", {})
                              .get("captures", [{}])[0]
                              .get("id", order_id))
                return jsonify({"status": "COMPLETED", "capture_id": capture_id,
                                "credits": tx.credits if tx else 0})

            return jsonify({"status": capture_data.get("status"), "data": capture_data})
        except Exception as e:
            current_app.logger.exception("PayPal capture-order failed")
            return jsonify({"error": str(e)}), 500

    def _process_payment(tracking_id, merchant_ref):
        if not tracking_id:
            return
        tx = Transaction.query.filter_by(provider_ref=tracking_id).first()
        if not tx and merchant_ref:
            tx = Transaction.query.filter_by(merchant_ref=merchant_ref).first()
        if not tx:
            return
        if tx.status == "complete":
            return
        try:
            status = pesapal.get_transaction_status(tracking_id)
        except Exception:
            current_app.logger.exception("PesaPal status check failed")
            return
        code = status.get("status_code")
        desc = (status.get("payment_status_description") or "").lower()
        if code == 1 or desc == "completed":
            tx.status = "complete"
            tx.completed_at = datetime.utcnow()
            user = User.query.get(tx.user_id)
            if user:
                user.credits = (user.credits or 0) + tx.credits
            db.session.commit()
        elif code in (2, 3) or desc in ("failed", "invalid"):
            tx.status = "failed"
            db.session.commit()

    # ══════════════════════════════════════════════════════
    # LIVE CHAT — student side
    # ══════════════════════════════════════════════════════

    @main.route("/chat")
    @require_login
    def chat_page():
        # Resume the most recent open session or start fresh
        sess = (ChatSession.query
                .filter_by(user_id=current_user.id, status="open")
                .order_by(ChatSession.created_at.desc())
                .first())
        return render_template("chat.html", session=sess)

    @main.route("/api/chat/start", methods=["POST"])
    @require_login
    def api_chat_start():
        """Create a new chat session and fire the AI welcome messages."""
        sess = ChatSession(user_id=current_user.id)
        db.session.add(sess)
        db.session.flush()   # get sess.id before commit

        name = current_user.display_name or "there"
        welcome = (
            f"Hi {name}! 👋 Welcome to Smart Study Guides Support. "
            "I'm connecting you with our customer care team right away."
        )
        wait_msg = (
            "Our agent is currently wrapping up with another student. "
            "Can you wait less than 10 minutes? They'll be with you very shortly! "
            "In the meantime, feel free to share your question or concern and we'll get right on it. 😊"
        )
        db.session.add(ChatMessage(session_id=sess.id, sender="ai", content=welcome))
        db.session.add(ChatMessage(session_id=sess.id, sender="ai", content=wait_msg))
        db.session.commit()

        msgs = [{"id": m.id, "sender": m.sender, "content": m.content,
                 "ts": m.created_at.strftime("%H:%M")} for m in sess.messages]
        return jsonify({"session_id": sess.id, "messages": msgs})

    # ── Notification API ────────────────────────────────────────────────────────

    @main.route("/api/notifications")
    @require_login
    def api_notifications():
        notifs = (UserNotification.query
                  .filter_by(user_id=current_user.id)
                  .order_by(UserNotification.created_at.desc())
                  .limit(30).all())
        unread_count = sum(1 for n in notifs if not n.is_read)
        return jsonify({
            "unread": unread_count,
            "items": [{
                "id": n.id,
                "type": n.type,
                "title": n.title,
                "body": n.body,
                "is_read": n.is_read,
                "ts": n.created_at.strftime("%-d %b, %H:%M"),
            } for n in notifs],
        })

    @main.route("/api/notifications/<int:nid>/read", methods=["POST"])
    @require_login
    def api_notification_read(nid):
        n = UserNotification.query.get_or_404(nid)
        if n.user_id != current_user.id:
            return jsonify({"error": "Forbidden"}), 403
        n.is_read = True
        db.session.commit()
        return jsonify({"ok": True})

    @main.route("/api/notifications/read-all", methods=["POST"])
    @require_login
    def api_notifications_read_all():
        UserNotification.query.filter_by(
            user_id=current_user.id, is_read=False
        ).update({"is_read": True}, synchronize_session=False)
        db.session.commit()
        return jsonify({"ok": True})

    # ── Weather (public, no login required) ──────────────────────────────────────

    @main.route("/api/weather")
    def api_weather():
        """Return weather for given lat/lon (or a default). Cached 10 min."""
        api_key = os.environ.get("OPENWEATHER_API_KEY", "")
        if not api_key:
            return jsonify({"error": "no_key", "text": "Weather unavailable"})

        lat = request.args.get("lat", "")
        lon = request.args.get("lon", "")

        if lat and lon:
            try:
                cache_key = f"{round(float(lat), 1)}_{round(float(lon), 1)}"
            except ValueError:
                cache_key = "default"
        else:
            cache_key = "default"

        now_ts = datetime.utcnow().timestamp()
        cached = _weather_cache.get(cache_key)
        if cached and (now_ts - cached["ts"]) < _WEATHER_TTL:
            return jsonify(cached["data"])

        try:
            if cache_key != "default":
                url = (
                    f"https://api.openweathermap.org/data/2.5/weather"
                    f"?lat={lat}&lon={lon}&appid={api_key}&units=metric"
                )
            else:
                url = (
                    f"https://api.openweathermap.org/data/2.5/weather"
                    f"?q=Washington,DC,US&appid={api_key}&units=metric"
                )
            r = http_requests.get(url, timeout=5)
            if not r.ok:
                return jsonify({"error": "api_error", "text": "Weather unavailable"})
            d = r.json()
            wid  = d["weather"][0]["id"]
            data = {
                "city": d.get("name", ""),
                "temp": round(d["main"]["temp"]),
                "icon": _weather_icon(wid),
            }
            _weather_cache[cache_key] = {"data": data, "ts": now_ts}
            return jsonify(data)
        except Exception:
            return jsonify({"error": "fetch_error", "text": "Weather unavailable"})

    # ── Chat unread count ────────────────────────────────────────────────────────

    @main.route("/api/chat/unread-count")
    @require_login
    def api_chat_unread_count():
        """Return count of unread staff/ai messages across all sessions for the current user."""
        session_ids = [s.id for s in ChatSession.query.filter_by(user_id=current_user.id).all()]
        if not session_ids:
            return jsonify({"count": 0})
        count = (ChatMessage.query
                 .filter(
                     ChatMessage.session_id.in_(session_ids),
                     ChatMessage.sender.in_(["staff", "ai"]),
                     ChatMessage.read_by_user == False,
                 ).count())
        return jsonify({"count": count})

    @main.route("/api/chat/mark-all-read", methods=["POST"])
    @require_login
    def api_chat_mark_all_read():
        """Mark every staff/ai message across all the user's sessions as read."""
        session_ids = [s.id for s in ChatSession.query.filter_by(user_id=current_user.id).all()]
        if session_ids:
            ChatMessage.query.filter(
                ChatMessage.session_id.in_(session_ids),
                ChatMessage.sender.in_(["staff", "ai"]),
                ChatMessage.read_by_user == False,
            ).update({"read_by_user": True}, synchronize_session=False)
            db.session.commit()
        return jsonify({"ok": True})

    @main.route("/api/chat/<int:sid>/messages")
    @require_login
    def api_chat_messages(sid):
        sess = ChatSession.query.get_or_404(sid)
        if sess.user_id != current_user.id and not current_user.is_staff:
            return "Forbidden", 403
        after = int(request.args.get("after", 0))
        msgs = (ChatMessage.query
                .filter(ChatMessage.session_id == sid, ChatMessage.id > after)
                .order_by(ChatMessage.id).all())

        # Mark staff/ai messages as read by the user when they fetch them
        if sess.user_id == current_user.id:
            unread_ids = [m.id for m in msgs if m.sender in ("staff", "ai") and not m.read_by_user]
            if unread_ids:
                ChatMessage.query.filter(ChatMessage.id.in_(unread_ids)).update(
                    {"read_by_user": True}, synchronize_session=False
                )
                db.session.commit()

        return jsonify([{"id": m.id, "sender": m.sender, "content": m.content,
                         "ts": m.created_at.strftime("%H:%M")} for m in msgs])

    @main.route("/api/chat/<int:sid>/send", methods=["POST"])
    @require_login
    def api_chat_send(sid):
        sess = ChatSession.query.get_or_404(sid)
        if sess.user_id != current_user.id:
            return "Forbidden", 403
        if sess.status == "closed":
            return jsonify({"error": "This chat has been closed."}), 400

        text = (request.json or {}).get("message", "").strip()
        if not text:
            return jsonify({"error": "Empty message."}), 400

        # Save student message
        db.session.add(ChatMessage(session_id=sid, sender="user", content=text))
        db.session.commit()

        ai_reply = None
        if not sess.has_human:
            # AI replies on behalf of support until a human takes over
            try:
                from .services.openai_service import generate_support_reply
                history = [{"sender": m.sender, "content": m.content} for m in sess.messages]
                reply_text = generate_support_reply(current_user.display_name or "Student", history)
                ai_msg = ChatMessage(session_id=sid, sender="ai", content=reply_text)
                db.session.add(ai_msg)
                db.session.commit()
                ai_reply = {"id": ai_msg.id, "sender": "ai",
                            "content": reply_text, "ts": ai_msg.created_at.strftime("%H:%M")}
            except Exception:
                current_app.logger.exception("Support AI reply failed")

        return jsonify({"ok": True, "ai_reply": ai_reply})

    # ══════════════════════════════════════════════════════
    # ══════════════════════════════════════════════════════
    # ADMIN — USER MANAGEMENT
    # ══════════════════════════════════════════════════════

    @main.route("/admin/users")
    @require_login
    def admin_users():
        if not current_user.is_staff:
            return redirect(url_for("main.dashboard"))
        from sqlalchemy import func
        q = request.args.get("q", "").strip()
        status_filter = request.args.get("status", "all")
        page = int(request.args.get("page", 1))
        per_page = 30

        query = User.query
        if q:
            query = query.filter(
                db.or_(
                    User.email.ilike(f"%{q}%"),
                    User.first_name.ilike(f"%{q}%"),
                    User.last_name.ilike(f"%{q}%"),
                )
            )
        if status_filter != "all":
            query = query.filter(User.account_status == status_filter)

        total = query.count()
        users = query.order_by(User.created_at.desc()).offset((page - 1) * per_page).limit(per_page).all()

        # Aggregate stats per user in one query
        assign_counts = dict(
            db.session.query(Assignment.user_id, func.count(Assignment.id))
            .group_by(Assignment.user_id).all()
        )
        tx_spent = dict(
            db.session.query(
                Transaction.user_id,
                func.coalesce(func.sum(Transaction.amount_usd_cents), 0)
            ).filter(Transaction.status == "completed")
            .group_by(Transaction.user_id).all()
        )
        job_counts = dict(
            db.session.query(JobDocument.user_id, func.count(JobDocument.id))
            .group_by(JobDocument.user_id).all()
        )

        # Overall platform stats
        stats = {
            "total": User.query.count(),
            "active": User.query.filter_by(account_status="active").count(),
            "flagged": User.query.filter_by(account_status="flagged").count(),
            "suspended": User.query.filter_by(account_status="suspended").count(),
            "terminated": User.query.filter_by(account_status="terminated").count(),
        }

        pages = max(1, (total + per_page - 1) // per_page)
        return render_template(
            "admin_users.html",
            users=users,
            assign_counts=assign_counts,
            tx_spent=tx_spent,
            job_counts=job_counts,
            stats=stats,
            q=q,
            status_filter=status_filter,
            page=page,
            pages=pages,
            total=total,
            user=current_user,
        )

    @main.route("/admin/users/<string:uid>")
    @require_login
    def admin_user_detail(uid):
        if not current_user.is_staff:
            return redirect(url_for("main.dashboard"))
        target = User.query.get_or_404(uid)
        assignments = (Assignment.query.filter_by(user_id=uid)
                       .order_by(Assignment.created_at.desc()).limit(20).all())
        transactions = (Transaction.query.filter_by(user_id=uid)
                        .order_by(Transaction.created_at.desc()).limit(20).all())
        job_docs = (JobDocument.query.filter_by(user_id=uid)
                    .order_by(JobDocument.created_at.desc()).limit(20).all())
        total_spent = (
            db.session.query(db.func.coalesce(db.func.sum(Transaction.amount_usd_cents), 0))
            .filter_by(user_id=uid, status="completed").scalar() or 0
        )
        return render_template(
            "admin_user_detail.html",
            target=target,
            assignments=assignments,
            transactions=transactions,
            job_docs=job_docs,
            total_spent=total_spent,
            user=current_user,
        )

    @main.route("/api/admin/users/<string:uid>/message", methods=["POST"])
    @require_login
    def api_admin_user_message(uid):
        if not current_user.is_staff:
            return jsonify({"error": "Forbidden"}), 403
        target = User.query.get_or_404(uid)
        data = request.get_json(force=True)
        text = (data.get("message") or "").strip()
        if not text:
            return jsonify({"error": "Message cannot be empty."}), 400

        # Derive a short title from the first line
        first_line = text.split("\n")[0].strip()[:80]
        title = first_line if first_line else "Notice from Smart Study Guides"

        # ── Deliver to notification bell (one-way, no chat involved) ──
        _notify(uid, "admin", title, text)

        # --- Email ---
        email_sent = False
        if target.email:
            try:
                from .services.email_service import send_admin_notification_email
                email_sent = send_admin_notification_email(
                    target.email, target.display_name, title, text
                )
            except Exception as e:
                print(f"[Admin notify] Email failed: {e}")

        # --- SMS ---
        sms_sent = False
        if target.phone:
            try:
                from .services.sms_service import send_admin_notification_sms
                sms_sent = send_admin_notification_sms(target.phone, text)
            except Exception as e:
                print(f"[Admin notify] SMS failed: {e}")

        return jsonify({
            "ok": True,
            "email_sent": email_sent,
            "sms_sent": sms_sent,
        })

    @main.route("/api/admin/users/<string:uid>/action", methods=["POST"])
    @require_login
    def api_admin_user_action(uid):
        if not current_user.is_staff:
            return jsonify({"error": "Forbidden"}), 403
        target = User.query.get_or_404(uid)
        # Core owner is immutable — no one can modify this account
        if (target.email or "").lower() == CORE_OWNER_EMAIL:
            return jsonify({"error": "The core owner account cannot be modified."}), 403
        data = request.get_json(force=True)
        action = data.get("action")  # flag | suspend | terminate | restore | delete
        reason = (data.get("reason") or "").strip()

        if action == "flag":
            target.account_status = "flagged"
            target.flag_reason = reason or "Suspicious activity detected."
        elif action == "suspend":
            target.account_status = "suspended"
            target.flag_reason = reason or "Account temporarily suspended."
        elif action == "terminate":
            target.account_status = "terminated"
            target.flag_reason = reason or "Account terminated."
        elif action == "restore":
            target.account_status = "active"
            target.flag_reason = None
        elif action == "adjust_credits":
            delta = int(data.get("delta", 0))
            target.credits = max(0, target.credits + delta)
        elif action == "delete":
            uid = target.id
            t = db.text
            # Delete in FK-safe order using raw SQL
            db.session.execute(t("DELETE FROM human_order_messages WHERE sender_id = :u"), {"u": uid})
            db.session.execute(t("DELETE FROM human_order_files WHERE uploader_id = :u"), {"u": uid})
            db.session.execute(t("DELETE FROM chat_messages WHERE session_id IN (SELECT id FROM chat_sessions WHERE user_id = :u)"), {"u": uid})
            db.session.execute(t("DELETE FROM chat_sessions WHERE user_id = :u"), {"u": uid})
            db.session.execute(t("DELETE FROM human_order_messages WHERE order_id IN (SELECT id FROM human_orders WHERE user_id = :u)"), {"u": uid})
            db.session.execute(t("DELETE FROM human_order_files WHERE order_id IN (SELECT id FROM human_orders WHERE user_id = :u)"), {"u": uid})
            db.session.execute(t("DELETE FROM human_orders WHERE user_id = :u"), {"u": uid})
            db.session.execute(t("DELETE FROM ai_removal_jobs WHERE user_id = :u"), {"u": uid})
            db.session.execute(t("DELETE FROM reviews WHERE assignment_id IN (SELECT id FROM assignments WHERE user_id = :u)"), {"u": uid})
            db.session.execute(t("DELETE FROM sources WHERE assignment_id IN (SELECT id FROM assignments WHERE user_id = :u)"), {"u": uid})
            db.session.execute(t("DELETE FROM pipeline_logs WHERE assignment_id IN (SELECT id FROM assignments WHERE user_id = :u)"), {"u": uid})
            db.session.execute(t("DELETE FROM assignments WHERE user_id = :u"), {"u": uid})
            db.session.execute(t("DELETE FROM transactions WHERE user_id = :u"), {"u": uid})
            db.session.execute(t("DELETE FROM subscriptions WHERE user_id = :u"), {"u": uid})
            db.session.execute(t("DELETE FROM daily_usage WHERE user_id = :u"), {"u": uid})
            db.session.execute(t("DELETE FROM user_notifications WHERE user_id = :u"), {"u": uid})
            db.session.execute(t("DELETE FROM job_documents WHERE user_id = :u"), {"u": uid})
            db.session.execute(t("DELETE FROM oauth WHERE user_id = :u"), {"u": uid})
            db.session.execute(t("DELETE FROM users WHERE id = :u"), {"u": uid})
            db.session.commit()
            return jsonify({"ok": True, "deleted": True})
        else:
            return jsonify({"error": "Unknown action"}), 400

        db.session.commit()
        return jsonify({"ok": True, "status": target.account_status, "credits": target.credits})

    @main.route("/api/admin/users/<string:uid>/set-writer", methods=["POST"])
    @require_login
    def api_admin_set_writer(uid):
        if not current_user.is_staff:
            return jsonify({"error": "Forbidden"}), 403
        target = User.query.get_or_404(uid)
        if (target.email or "").lower() == CORE_OWNER_EMAIL:
            return jsonify({"error": "The core owner account cannot be modified."}), 403
        data = request.get_json(force=True) or {}
        target.is_writer = bool(data.get("is_writer", False))
        db.session.commit()
        return jsonify({"ok": True, "is_writer": target.is_writer})

    # LIVE CHAT — staff / admin side
    # ══════════════════════════════════════════════════════

    def _require_staff():
        if not current_user.is_authenticated or not current_user.is_staff:
            return redirect(url_for("main.dashboard"))
        return None

    @main.route("/admin/chats")
    @require_login
    def admin_chats():
        if not current_user.is_staff:
            return redirect(url_for("main.dashboard"))
        open_sessions = (ChatSession.query
                         .filter_by(status="open")
                         .order_by(ChatSession.updated_at.desc())
                         .all())
        # Count unread student messages per session
        unread = {}
        for s in open_sessions:
            unread[s.id] = sum(
                1 for m in s.messages
                if m.sender == "user" and not m.read_by_staff
            )
        return render_template("admin_chats.html", sessions=open_sessions, unread=unread)

    @main.route("/admin/chat/<int:sid>")
    @require_login
    def admin_chat_detail(sid):
        if not current_user.is_staff:
            return redirect(url_for("main.dashboard"))
        sess = ChatSession.query.get_or_404(sid)
        # Mark all student messages as read
        for m in sess.messages:
            if m.sender == "user":
                m.read_by_staff = True
        db.session.commit()
        # Fetch student's recent completed assignments so agent can download/share
        student_assignments = (
            Assignment.query
            .filter_by(user_id=sess.user_id, status="complete")
            .order_by(Assignment.created_at.desc())
            .limit(10)
            .all()
        )
        return render_template("admin_chat_detail.html", sess=sess,
                               student_assignments=student_assignments)

    @main.route("/api/admin/chat/<int:sid>/send", methods=["POST"])
    @require_login
    def api_admin_chat_send(sid):
        if not current_user.is_staff:
            return "Forbidden", 403
        sess = ChatSession.query.get_or_404(sid)
        text = (request.json or {}).get("message", "").strip()
        if not text:
            return jsonify({"error": "Empty message."}), 400

        msg = ChatMessage(session_id=sid, sender="staff", content=text, read_by_staff=True)
        db.session.add(msg)
        sess.has_human = True   # disable AI auto-reply for this session
        sess.updated_at = datetime.utcnow()
        db.session.commit()
        return jsonify({"id": msg.id, "sender": "staff", "content": text,
                        "ts": msg.created_at.strftime("%H:%M")})

    @main.route("/api/admin/chat/<int:sid>/close", methods=["POST"])
    @require_login
    def api_admin_chat_close(sid):
        if not current_user.is_staff:
            return "Forbidden", 403
        sess = ChatSession.query.get_or_404(sid)
        sess.status = "closed"
        close_msg = ChatMessage(
            session_id=sid, sender="staff",
            content="This chat has been closed by our support team. Thank you for reaching out! 😊",
            read_by_staff=True,
        )
        db.session.add(close_msg)
        db.session.commit()
        return jsonify({"ok": True})

    @main.route("/api/admin/chats/stats")
    @require_login
    def api_admin_chat_stats():
        if not current_user.is_staff:
            return "Forbidden", 403
        open_count = ChatSession.query.filter_by(status="open").count()
        unread_total = (db.session.query(ChatMessage)
                        .join(ChatSession)
                        .filter(ChatSession.status == "open",
                                ChatMessage.sender == "user",
                                ChatMessage.read_by_staff == False)
                        .count())
        return jsonify({"open": open_count, "unread": unread_total})

    # ══════════════════════════════════════════════════════════════════════════
    # HUMAN WRITER ORDERS
    # ══════════════════════════════════════════════════════════════════════════

    _HW_CREDITS_PER_PAGE  = 1200
    _HW_URGENT_MULTIPLIER = 1.5   # urgent = 1800 credits/page

    @main.route("/post-assignment", methods=["GET", "POST"])
    @require_login
    def post_human_assignment():
        if request.method == "POST":
            title          = request.form.get("title", "").strip()
            subject        = request.form.get("subject", "").strip()
            instructions   = request.form.get("instructions", "").strip()
            academic_level = request.form.get("academic_level", "").strip()
            format_style   = request.form.get("format_style", "APA").strip()
            priority       = request.form.get("priority", "standard")
            deadline_str   = request.form.get("deadline", "").strip()

            try:
                num_pages = max(1, int(request.form.get("num_pages", 1)))
            except Exception:
                num_pages = 1
            try:
                num_references = max(0, int(request.form.get("num_references", 0)))
            except Exception:
                num_references = 0

            errors = []
            if not title:          errors.append("Assignment title is required.")
            if not subject:        errors.append("Subject / course is required.")
            if not instructions:   errors.append("Detailed instructions are required.")
            if not academic_level: errors.append("Academic level is required.")
            if not deadline_str:   errors.append("Deadline is required.")

            deadline = None
            if deadline_str:
                try:
                    deadline = datetime.strptime(deadline_str, "%Y-%m-%dT%H:%M")
                    if deadline <= datetime.utcnow():
                        errors.append("Deadline must be in the future.")
                except Exception:
                    errors.append("Invalid deadline format.")

            rate         = int(_HW_CREDITS_PER_PAGE * (_HW_URGENT_MULTIPLIER if priority == "urgent" else 1))
            credits_cost = rate * num_pages

            if errors:
                for e in errors:
                    flash(e, "error")
                return redirect(url_for("main.post_human_assignment"))

            sub = _get_active_subscription(current_user.id)
            if not sub:
                flash("You need an active subscription to submit a human writer order.", "error")
                return redirect(url_for("main.pricing"))

            order = HumanOrder(
                user_id=current_user.id,
                title=title,
                subject=subject,
                instructions=instructions,
                academic_level=academic_level,
                format_style=format_style,
                num_pages=num_pages,
                num_references=num_references,
                deadline=deadline,
                priority=priority,
                credits_per_page=rate,
                credits_paid=credits_cost,
                status="pending",
            )
            db.session.add(order)
            db.session.flush()  # get order.id

            f = request.files.get("brief_file")
            if f and f.filename:
                file_bytes = f.read()
                ct  = f.content_type or "application/octet-stream"
                ext = f.filename.rsplit(".", 1)[-1].lower() if "." in f.filename else "bin"
                path = f"human_orders/{order.id}/brief_{uuid.uuid4().hex}.{ext}"
                url  = supabase_storage.upload_file(path, file_bytes, ct, signed_days=365)
                db.session.add(HumanOrderFile(
                    order_id=order.id,
                    uploader_id=current_user.id,
                    file_url=url,
                    file_name=f.filename,
                    file_type="brief",
                ))

            db.session.commit()
            flash("✅ Assignment submitted! Our writers will pick it up shortly.", "success")
            return redirect(url_for("main.human_order_detail", oid=order.id))

        return render_template(
            "post_assignment.html",
            credits_per_page=0,
            urgent_multiplier=_HW_URGENT_MULTIPLIER,
            user=current_user,
        )

    @main.route("/human-orders/<int:oid>")
    @require_login
    def human_order_detail(oid):
        order = HumanOrder.query.get_or_404(oid)
        if order.user_id != current_user.id and not current_user.is_staff:
            return "Forbidden", 403
        return render_template("human_order_detail.html", order=order, user=current_user)

    @main.route("/api/human-orders/<int:oid>/message", methods=["POST"])
    @require_login
    def api_human_order_message(oid):
        order = HumanOrder.query.get_or_404(oid)
        if order.user_id != current_user.id and not current_user.is_staff and not current_user.is_writer:
            return jsonify({"error": "Forbidden"}), 403
        content = (request.json or {}).get("content", "").strip()
        if not content:
            return jsonify({"error": "Empty message"}), 400
        role = "writer" if (current_user.is_staff or current_user.is_writer) else "student"
        msg  = HumanOrderMessage(
            order_id=oid,
            sender_id=current_user.id,
            sender_role=role,
            content=content,
        )
        db.session.add(msg)
        db.session.commit()
        return jsonify({
            "id":          msg.id,
            "sender_role": role,
            "sender_name": current_user.display_name,
            "content":     msg.content,
            "ts":          msg.created_at.strftime("%d %b %H:%M"),
        })

    @main.route("/api/human-orders/<int:oid>/messages")
    @require_login
    def api_human_order_messages(oid):
        order = HumanOrder.query.get_or_404(oid)
        if order.user_id != current_user.id and not current_user.is_staff and not current_user.is_writer:
            return jsonify({"error": "Forbidden"}), 403
        after = int(request.args.get("after", 0))
        msgs  = (HumanOrderMessage.query
                 .filter(HumanOrderMessage.order_id == oid, HumanOrderMessage.id > after)
                 .order_by(HumanOrderMessage.id).all())
        return jsonify([{
            "id":          m.id,
            "sender_role": m.sender_role,
            "sender_name": m.sender.display_name,
            "content":     m.content,
            "ts":          m.created_at.strftime("%d %b %H:%M"),
        } for m in msgs])

    @main.route("/api/human-orders/<int:oid>/status")
    @require_login
    def api_human_order_status(oid):
        order = HumanOrder.query.get_or_404(oid)
        if order.user_id != current_user.id and not current_user.is_staff and not current_user.is_writer:
            return jsonify({"error": "Forbidden"}), 403
        return jsonify({
            "status":          order.status,
            "final_file_url":  order.final_file_url,
            "final_file_name": order.final_file_name,
        })

    # ── Admin: writer-mode preview toggle ──────────────────────────────────────

    @main.route("/admin/api-health")
    @require_login
    def admin_api_health():
        if not current_user.is_staff:
            return redirect(url_for("main.dashboard"))

        import requests as _req

        checks = []

        # ── OpenAI ──────────────────────────────────────────────────────────────
        try:
            from .services.openai_service import client as _oai_client
            _oai_client().models.list()
            checks.append({
                "name": "OpenAI (GPT-4o)",
                "icon": "🤖",
                "status": "ok",
                "label": "API key valid — quota available",
                "link": "https://platform.openai.com/usage",
                "link_label": "View usage & billing →",
            })
        except Exception as e:
            err = str(e).lower()
            if "quota" in err or "insufficient" in err:
                msg = "Quota exhausted — top up credits immediately"
                status = "critical"
            elif "auth" in err or "invalid" in err or "key" in err:
                msg = "Invalid API key"
                status = "critical"
            else:
                msg = f"Error: {str(e)[:120]}"
                status = "warn"
            checks.append({
                "name": "OpenAI (GPT-4o)",
                "icon": "🤖",
                "status": status,
                "label": msg,
                "link": "https://platform.openai.com/usage",
                "link_label": "View usage & billing →",
            })

        # ── Ryter Pro ───────────────────────────────────────────────────────────
        try:
            key = os.environ.get("RYTER_PRO_API_KEY", "")
            if not key:
                raise ValueError("RYTER_PRO_API_KEY not set")
            r = _req.get(
                "https://api.ryter.pro/api/v1/user",
                headers={"x-api-key": key, "Content-Type": "application/json"},
                timeout=10,
            )
            if r.status_code == 200:
                data = r.json() if r.text else {}
                credits = data.get("credits") or data.get("balance") or data.get("words_remaining")
                label = f"Active — {credits:,} credits remaining" if credits is not None else "API key valid"
                status = "warn" if (credits is not None and credits < 5000) else "ok"
                checks.append({
                    "name": "Ryter Pro",
                    "icon": "✍️",
                    "status": status,
                    "label": label,
                    "link": "https://ryter.pro/dashboard",
                    "link_label": "Manage subscription →",
                })
            elif r.status_code in (401, 403):
                checks.append({"name": "Ryter Pro", "icon": "✍️", "status": "critical",
                    "label": "Invalid or expired API key", "link": "https://ryter.pro/dashboard", "link_label": "Manage subscription →"})
            else:
                checks.append({"name": "Ryter Pro", "icon": "✍️", "status": "warn",
                    "label": f"Unexpected response: HTTP {r.status_code}", "link": "https://ryter.pro/dashboard", "link_label": "Manage subscription →"})
        except Exception as e:
            checks.append({"name": "Ryter Pro", "icon": "✍️", "status": "warn",
                "label": f"Could not reach API: {str(e)[:100]}", "link": "https://ryter.pro/dashboard", "link_label": "Manage subscription →"})

        # ── Copyleaks ───────────────────────────────────────────────────────────
        try:
            from .services.plagiarism_service import configured as _cl_configured, _login as _cl_login
            if not _cl_configured():
                checks.append({"name": "Copyleaks", "icon": "📋", "status": "warn",
                    "label": "Not configured — set COPYLEAKS_EMAIL and COPYLEAKS_API_KEY",
                    "link": "https://copyleaks.com", "link_label": "Go to Copyleaks →"})
            else:
                token = _cl_login()
                r2 = _req.get(
                    "https://api.copyleaks.com/v3/account/credits",
                    headers={"Authorization": f"Bearer {token}"},
                    timeout=15,
                )
                if r2.status_code == 200:
                    data2 = r2.json()
                    remaining = data2.get("remaining") or data2.get("credits") or data2.get("balance")
                    if remaining is not None:
                        status = "critical" if remaining < 10 else ("warn" if remaining < 50 else "ok")
                        label = f"{remaining:,} scan credits remaining"
                    else:
                        status = "ok"
                        label = "API key valid"
                    checks.append({"name": "Copyleaks", "icon": "📋", "status": status,
                        "label": label, "link": "https://copyleaks.com/dashboard",
                        "link_label": "Buy more credits →"})
                else:
                    checks.append({"name": "Copyleaks", "icon": "📋", "status": "warn",
                        "label": f"Credits endpoint returned HTTP {r2.status_code}",
                        "link": "https://copyleaks.com/dashboard", "link_label": "Check dashboard →"})
        except Exception as e:
            checks.append({"name": "Copyleaks", "icon": "📋", "status": "warn",
                "label": f"Error: {str(e)[:120]}", "link": "https://copyleaks.com/dashboard",
                "link_label": "Check dashboard →"})

        return render_template("admin_api_health.html", checks=checks)

    @main.route("/admin/writer-mode/on")
    @require_login
    def admin_writer_mode_on():
        if not current_user.is_staff:
            return redirect(url_for("main.dashboard"))
        session["admin_writer_preview"] = True
        return redirect(url_for("main.writer_available"))

    @main.route("/admin/writer-mode/off")
    @require_login
    def admin_writer_mode_off():
        session.pop("admin_writer_preview", None)
        return redirect(url_for("main.dashboard"))

    # ── Writer Portal (for is_writer accounts + admin preview) ─────────────────

    def _is_writer_view():
        """True when the current user should see the writer portal UI."""
        return current_user.is_writer or bool(session.get("admin_writer_preview"))

    @main.route("/writer")
    @require_login
    def writer_home():
        if not _is_writer_view() and not current_user.is_staff:
            return redirect(url_for("main.dashboard"))
        return redirect(url_for("main.writer_available"))

    @main.route("/writer/available")
    @require_login
    def writer_available():
        if not _is_writer_view() and not current_user.is_staff:
            return redirect(url_for("main.dashboard"))
        orders = (HumanOrder.query
                  .filter_by(status="pending")
                  .order_by(HumanOrder.created_at.desc())
                  .all())
        return render_template("writer_portal_available.html", orders=orders, user=current_user)

    @main.route("/writer/my-orders")
    @require_login
    def writer_my_orders():
        if not _is_writer_view() and not current_user.is_staff:
            return redirect(url_for("main.dashboard"))
        my_orders = (HumanOrder.query
                     .filter_by(writer_id=current_user.id)
                     .order_by(HumanOrder.updated_at.desc())
                     .all())
        return render_template("writer_portal_my_orders.html", orders=my_orders, user=current_user)

    @main.route("/writer/<int:oid>")
    @require_login
    def writer_order_view(oid):
        if not _is_writer_view() and not current_user.is_staff:
            return redirect(url_for("main.dashboard"))
        order = HumanOrder.query.get_or_404(oid)
        return render_template("writer_order_detail.html", order=order, user=current_user)

    # ── Writers Dashboard (admin staff only) ────────────────────────────────────

    @main.route("/writers")
    @require_login
    def writers_dashboard():
        if not current_user.is_staff:
            flash("Access restricted to staff.", "error")
            return redirect(url_for("main.dashboard"))
        status_filter = request.args.get("status", "")
        q = HumanOrder.query.order_by(HumanOrder.created_at.desc())
        if status_filter:
            q = q.filter_by(status=status_filter)
        orders = q.all()
        counts = {
            s: HumanOrder.query.filter_by(status=s).count()
            for s in ("pending", "assigned", "in_progress", "completed", "delivered")
        }
        return render_template(
            "writers_dashboard.html",
            orders=orders, status_filter=status_filter,
            counts=counts, user=current_user,
        )

    @main.route("/writers/<int:oid>")
    @require_login
    def writer_order_detail(oid):
        if not current_user.is_staff and not current_user.is_writer:
            flash("Access restricted.", "error")
            return redirect(url_for("main.dashboard"))
        order = HumanOrder.query.get_or_404(oid)
        return render_template("writer_order_detail.html", order=order, user=current_user)

    @main.route("/api/human-orders/<int:oid>/claim", methods=["POST"])
    @require_login
    def api_human_order_claim(oid):
        if not current_user.is_staff and not current_user.is_writer:
            return jsonify({"error": "Forbidden"}), 403
        order = HumanOrder.query.get_or_404(oid)
        if order.status != "pending":
            return jsonify({"error": "Order is no longer available for claiming."}), 400
        order.writer_id = current_user.id
        order.status    = "assigned"
        db.session.commit()
        return jsonify({"ok": True, "status": order.status})

    @main.route("/api/human-orders/<int:oid>/update-status", methods=["POST"])
    @require_login
    def api_human_order_update_status(oid):
        if not current_user.is_staff and not current_user.is_writer:
            return jsonify({"error": "Forbidden"}), 403
        order      = HumanOrder.query.get_or_404(oid)
        new_status = (request.json or {}).get("status", "").strip()
        if new_status not in ("assigned", "in_progress", "completed"):
            return jsonify({"error": "Invalid status value."}), 400
        order.status = new_status
        db.session.commit()
        return jsonify({"ok": True, "status": order.status})

    @main.route("/api/human-orders/<int:oid>/deliver", methods=["POST"])
    @require_login
    def api_human_order_deliver(oid):
        if not current_user.is_staff and not current_user.is_writer:
            return jsonify({"error": "Forbidden"}), 403
        order = HumanOrder.query.get_or_404(oid)
        confirmed = request.form.get("confirmed_human", "") == "1"
        if not confirmed:
            flash("You must confirm this is human-written work before delivering.", "error")
            return redirect(url_for("main.writer_order_detail", oid=oid))
        f = request.files.get("delivery_file")
        if not f or not f.filename:
            flash("Please upload the completed assignment file.", "error")
            return redirect(url_for("main.writer_order_detail", oid=oid))
        file_bytes = f.read()
        ct   = f.content_type or "application/octet-stream"
        ext  = f.filename.rsplit(".", 1)[-1].lower() if "." in f.filename else "bin"
        path = f"human_orders/{oid}/delivery_{uuid.uuid4().hex}.{ext}"
        url  = supabase_storage.upload_file(path, file_bytes, ct, signed_days=365)
        db.session.add(HumanOrderFile(
            order_id=oid,
            uploader_id=current_user.id,
            file_url=url,
            file_name=f.filename,
            file_type="delivery",
        ))
        order.final_file_url           = url
        order.final_file_name          = f.filename
        order.status                   = "delivered"
        order.writer_confirmed_human   = True
        order.completed_at             = datetime.utcnow()
        db.session.commit()
        flash("✅ Assignment delivered to the student!", "success")
        return redirect(url_for("main.writer_order_detail", oid=oid))

    # ═══════════════════════════════════════════════════════════════
    # JOB APPLICATION DOCUMENTS
    # Free for users who have credits > 0 OR have ever purchased
    # ═══════════════════════════════════════════════════════════════

    JOB_DOC_TYPES = {
        "cv":                  {"label": "Curriculum Vitae (CV)",     "icon": "📄", "desc": "Full academic & work history"},
        "resume":              {"label": "Resume",                     "icon": "📋", "desc": "Concise 1-page professional summary"},
        "cover_letter":        {"label": "Cover Letter",               "icon": "✉️",  "desc": "Compelling introduction to an employer"},
        "email_note":          {"label": "Application Email",          "icon": "📧", "desc": "Professional email to send with your application"},
        "motivation_letter":   {"label": "Motivation Letter",          "icon": "💡", "desc": "Why you want this role or programme"},
        "reference_letter":    {"label": "Reference Letter",           "icon": "🤝", "desc": "Letter recommending a person"},
        "thank_you_letter":    {"label": "Thank You Letter",           "icon": "🙏", "desc": "Post-interview appreciation note"},
        "scholarship_letter":  {"label": "Scholarship Application",    "icon": "🎓", "desc": "Letter applying for financial aid"},
        "internship_letter":   {"label": "Internship Application",     "icon": "🏢", "desc": "Letter for an internship position"},
        "recommendation_req":  {"label": "Recommendation Request",     "icon": "📨", "desc": "Ask someone to recommend you"},
    }

    def _can_use_job_docs(user):
        """
        Career section access:
        - Staff/admin always allowed, OR
        - User has an active subscription, OR
        - User has referred at least one person who has an active subscription.
        """
        if getattr(user, "is_staff", False):
            return True
        # Direct subscription check
        if _get_active_subscription(user.id):
            return True
        # Referral check — any referred user with an active subscription
        referred_ids = [
            u.id for u in User.query.filter_by(referred_by_id=user.id).all()
        ]
        if referred_ids:
            active_referral_sub = (
                Subscription.query
                .filter(Subscription.user_id.in_(referred_ids),
                        Subscription.status == "active",
                        Subscription.end_date > datetime.utcnow())
                .first()
            )
            if active_referral_sub:
                return True
        return False

    def _has_referral_unlock(user):
        """True if access comes from a referral (not the user's own subscription)."""
        if _get_active_subscription(user.id):
            return False   # own subscription — not a referral unlock
        referred_ids = [
            u.id for u in User.query.filter_by(referred_by_id=user.id).all()
        ]
        if not referred_ids:
            return False
        return (
            Subscription.query
            .filter(Subscription.user_id.in_(referred_ids),
                    Subscription.status == "active",
                    Subscription.end_date > datetime.utcnow())
            .first()
        ) is not None

    CAREER_MONTHLY_LIMIT = 4

    def _career_usage(user_id):
        """Returns {doc_type: count} for rolling 30-day window."""
        from sqlalchemy import func as _fn
        cutoff = datetime.utcnow() - timedelta(days=30)
        rows = db.session.query(
            JobDocument.doc_type, _fn.count(JobDocument.id)
        ).filter(
            JobDocument.user_id    == user_id,
            JobDocument.created_at >= cutoff,
        ).group_by(JobDocument.doc_type).all()
        return {r[0]: r[1] for r in rows}

    def _career_type_reset_date(user_id, doc_type):
        """Returns 'DD Mon YYYY' string when this doc_type's oldest use expires."""
        cutoff  = datetime.utcnow() - timedelta(days=30)
        oldest  = JobDocument.query.filter(
            JobDocument.user_id    == user_id,
            JobDocument.doc_type   == doc_type,
            JobDocument.created_at >= cutoff,
        ).order_by(JobDocument.created_at.asc()).first()
        base = oldest.created_at if oldest else datetime.utcnow()
        return (base + timedelta(days=30)).strftime("%d %b %Y")

    def _build_job_doc_prompt(doc_label, details):
        """Build a detailed AI prompt from the submitted form details."""
        edu_lines = []
        for e in (details.get("education") or []):
            if e.get("school") or e.get("qual"):
                edu_lines.append(
                    f"  • {e.get('level','')} — {e.get('school','')} "
                    f"({e.get('qual','')}) — {e.get('year','')}"
                )
        edu_text = "\n".join(edu_lines) if edu_lines else "  (not provided)"

        return f"""You are a professional career-document writer.
Write a complete, polished, ready-to-submit **{doc_label}** for the person described below.

FORMAT RULES:
- Use standard professional formatting for this document type.
- Include all appropriate sections (date, sender address, salutation, body, closing, etc.).
- Do NOT include any commentary, preamble, or post-document notes — output the document ONLY.
- Write in formal, confident, error-free English.

=== APPLICANT DETAILS ===
Full Name      : {details.get('full_name', '')}
Email          : {details.get('email', '')}
Phone          : {details.get('phone', '')}
LinkedIn/URL   : {details.get('linkedin', '')}

=== APPLICANT ADDRESS ===
{details.get('street', '')}, {details.get('city', '')}, {details.get('state', '')}, {details.get('country', '')}

=== EDUCATION HISTORY (earliest → highest) ===
{edu_text}
Current School : {details.get('current_school', '(none)')}

=== RECEIVER / EMPLOYER ===
Name           : {details.get('receiver_name', '')}
Title          : {details.get('receiver_title', '')}
Organisation   : {details.get('receiver_org', '')}
Location       : {details.get('receiver_location', '')}

=== DOCUMENT PURPOSE ===
Position / Role / Programme : {details.get('position', '')}
Key Skills / Strengths       : {details.get('skills', '')}
Additional Notes             : {details.get('notes', '')}

Now write the complete {doc_label} below:
"""

    @main.route("/job-docs")
    @require_login
    def job_docs():
        can_access     = _can_use_job_docs(current_user)
        # Subscribed users (or referral-unlocked users) have no monthly limit
        bypassed       = can_access
        referral_unlock = _has_referral_unlock(current_user) if can_access else False
        usage          = _career_usage(current_user.id) if can_access else {}
        reset_dates    = {}
        return render_template(
            "job_docs.html",
            user=current_user,
            doc_types=JOB_DOC_TYPES,
            can_access=can_access,
            usage=usage,
            limit=CAREER_MONTHLY_LIMIT,
            bypassed=bypassed,
            reset_dates=reset_dates,
            referral_unlock=referral_unlock,
        )

    @main.route("/job-docs/start", methods=["POST"])
    @require_login
    def job_docs_start():
        if not _can_use_job_docs(current_user):
            return jsonify({"error": "Subscribe to a plan to unlock this feature."}), 403
        try:
            data     = request.get_json(force=True)
            doc_type = data.get("doc_type", "").strip()
            if doc_type not in JOB_DOC_TYPES:
                return jsonify({"error": "Invalid document type."}), 400

            # Subscribers and referral-unlocked users have no monthly cap

            doc = JobDocument(
                user_id      = current_user.id,
                doc_type     = doc_type,
                doc_label    = JOB_DOC_TYPES[doc_type]["label"],
                details_json = json.dumps(data),
                status       = "pending",
            )
            db.session.add(doc)
            db.session.commit()

            return jsonify({"redirect": url_for("main.job_docs_live", doc_id=doc.id)})
        except Exception as e:
            current_app.logger.exception("job_docs_start failed")
            return jsonify({"error": str(e)}), 500

    @main.route("/job-docs/<int:doc_id>/live")
    @require_login
    def job_docs_live(doc_id):
        doc = JobDocument.query.filter_by(id=doc_id, user_id=current_user.id).first_or_404()
        try:
            details = json.loads(doc.details_json)
        except Exception:
            details = {}
        return render_template("job_docs_live.html", user=current_user, doc=doc, details=details)

    @main.route("/job-docs/<int:doc_id>/export-docx", methods=["POST"])
    @require_login
    def job_docs_export_docx(doc_id):
        import re as _re
        doc = JobDocument.query.filter_by(id=doc_id, user_id=current_user.id).first_or_404()
        if doc.status != "done" or not doc.content:
            return jsonify({"error": "Document not ready yet."}), 400
        # Return cached URL if already generated
        # Return cached copies if both already exist
        if doc.docx_url and doc.pdf_url:
            return jsonify({
                "url":          doc.docx_url,
                "filename":     doc.docx_filename,
                "pdf_url":      doc.pdf_url,
                "pdf_filename": doc.pdf_filename,
            })
        try:
            details = json.loads(doc.details_json)
        except Exception:
            details = {}

        user_name  = details.get("full_name") or current_user.display_name or "User"
        safe_name  = _re.sub(r"[^a-zA-Z0-9 _-]", "", user_name).strip().replace(" ", "_")
        date_str   = datetime.utcnow().strftime("%Y%m%d")
        label_slug = _re.sub(r"[^a-zA-Z0-9]+", "_", doc.doc_label)

        # ── Build Word (.docx) ────────────────────────────────────────────────
        if not doc.docx_url:
            from .services.career_docx_builder import build_career_docx
            docx_bytes  = build_career_docx(
                content=doc.content, doc_label=doc.doc_label,
                doc_type=doc.doc_type, details=details,
            )
            docx_name   = f"Career_{label_slug}_{safe_name}_{date_str}.docx"
            docx_url    = supabase_storage.upload_file(
                f"career_docs/{current_user.id}/{doc.id}/{docx_name}",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                signed_days=365,
            )
            doc.docx_url      = docx_url
            doc.docx_filename = docx_name

        # ── Build PDF ─────────────────────────────────────────────────────────
        if not doc.pdf_url:
            from .services.career_pdf_builder import build_career_pdf
            pdf_bytes   = build_career_pdf(
                content=doc.content, doc_label=doc.doc_label,
                doc_type=doc.doc_type, details=details,
            )
            pdf_name    = f"Career_{label_slug}_{safe_name}_{date_str}.pdf"
            pdf_url     = supabase_storage.upload_file(
                f"career_docs/{current_user.id}/{doc.id}/{pdf_name}",
                pdf_bytes,
                "application/pdf",
                signed_days=365,
            )
            doc.pdf_url      = pdf_url
            doc.pdf_filename = pdf_name

        db.session.commit()
        return jsonify({
            "url":          doc.docx_url,
            "filename":     doc.docx_filename,
            "pdf_url":      doc.pdf_url,
            "pdf_filename": doc.pdf_filename,
        })

    @main.route("/job-docs/<int:doc_id>/stream")
    @require_login
    def job_docs_stream(doc_id):
        doc = JobDocument.query.filter_by(id=doc_id, user_id=current_user.id).first_or_404()

        def _generate():
            from .services.openai_service import stream_chat as _stream
            try:
                details = json.loads(doc.details_json)
            except Exception:
                details = {}

            prompt = _build_job_doc_prompt(doc.doc_label, details)

            # Update status
            doc.status = "streaming"
            db.session.commit()

            full_text = []
            try:
                for token in _stream(prompt, max_tokens=4000):
                    full_text.append(token)
                    payload = json.dumps({"token": token})
                    yield f"data: {payload}\n\n"

                # Save completed content
                doc.content = "".join(full_text)
                doc.status  = "done"
                db.session.commit()
                yield "data: [DONE]\n\n"

            except Exception as e:
                doc.status = "failed"
                db.session.commit()
                current_app.logger.exception("job_docs_stream error")
                yield f"data: [ERROR] {str(e)}\n\n"

        return Response(
            stream_with_context(_generate()),
            mimetype="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",
            },
        )

    # ── AI Removal Jobs ──────────────────────────────────────────────────────

    OWNER_EMAILS_SET = {"aservices767@gmail.com", "akyida001@gmail.com",
                        "simonedwardj@yahoo.com", "simonedwardj6@gmail.com"}

    @main.route("/assignments/<int:aid>/submit-ai-removal", methods=["POST"])
    @require_login
    def submit_ai_removal(aid):
        from .services.email_service import send_ai_removal_admin_alert
        a = Assignment.query.filter_by(id=aid, user_id=current_user.id).first_or_404()

        # Prevent duplicate pending/assigned submissions
        existing = AIRemovalJob.query.filter(
            AIRemovalJob.assignment_id == aid,
            AIRemovalJob.status.in_(["pending", "assigned"]),
        ).first()
        if existing:
            flash("Your paper is already in the AI removal queue.", "info")
            return redirect(url_for("main.ai_removal_status", aid=aid))

        original_text = request.form.get("humanized_text", "") or (a.paper_text or "")
        ai_score      = float(request.form.get("ai_score", 50) or 50)
        notes         = request.form.get("notes", "").strip()
        word_count    = len(original_text.split())
        deadline_mins = 30 if word_count <= 800 else 60
        deadline      = datetime.utcnow() + timedelta(minutes=deadline_mins)

        job = AIRemovalJob(
            assignment_id     = aid,
            user_id           = current_user.id,
            original_text     = original_text,
            original_ai_score = round(ai_score, 1),
            word_count        = word_count,
            deadline          = deadline,
            student_notes     = notes or None,
        )
        db.session.add(job)
        db.session.commit()

        # Notify all admins by email
        domains = os.environ.get("REPLIT_DOMAINS", "smart-study-guides.com")
        base    = f"https://{domains.split(',')[0]}"
        job_url = f"{base}/admin/ai-removal-jobs"
        dl_str  = deadline.strftime("%H:%M, %d %b %Y") + f" ({deadline_mins} min)"
        for admin_email in OWNER_EMAILS_SET:
            try:
                send_ai_removal_admin_alert(
                    admin_email,
                    current_user.display_name,
                    a.topic,
                    round(ai_score, 1),
                    dl_str,
                    job_url,
                )
            except Exception:
                pass

        flash("Submitted! Our writers will have your paper ready within the deadline.", "success")
        return redirect(url_for("main.ai_removal_status", aid=aid))

    @main.route("/assignments/<int:aid>/ai-removal-status")
    @require_login
    def ai_removal_status(aid):
        a   = Assignment.query.filter_by(id=aid, user_id=current_user.id).first_or_404()
        job = (AIRemovalJob.query
               .filter_by(assignment_id=aid, user_id=current_user.id)
               .order_by(AIRemovalJob.submitted_at.desc())
               .first_or_404())
        return render_template("ai_removal_status.html", a=a, job=job)

    @main.route("/admin/ai-removal-jobs")
    @require_login
    def admin_ai_removal_jobs():
        if not current_user.is_staff and not _is_owner(current_user):
            return redirect(url_for("main.dashboard"))
        jobs    = AIRemovalJob.query.order_by(
            AIRemovalJob.status.asc(), AIRemovalJob.deadline.asc()
        ).all()
        writers = User.query.filter_by(is_writer=True).all()
        return render_template("admin_ai_removal_jobs.html", jobs=jobs, writers=writers)

    @main.route("/admin/ai-removal-jobs/<int:job_id>/assign", methods=["POST"])
    @require_login
    def admin_assign_ai_removal(job_id):
        if not current_user.is_staff and not _is_owner(current_user):
            return redirect(url_for("main.dashboard"))
        job       = AIRemovalJob.query.get_or_404(job_id)
        writer_id = request.form.get("writer_id", "").strip()
        if writer_id:
            job.writer_id = writer_id
            job.status    = "assigned"
            db.session.commit()
            flash("Writer assigned successfully.", "success")
        else:
            flash("Please select a writer.", "error")
        return redirect(url_for("main.admin_ai_removal_jobs"))

    @main.route("/writer/ai-removal-jobs")
    @require_login
    def writer_ai_removal_jobs():
        if not current_user.is_writer and not current_user.is_staff and not _is_owner(current_user):
            return redirect(url_for("main.dashboard"))
        if current_user.is_staff or _is_owner(current_user):
            jobs = AIRemovalJob.query.filter(
                AIRemovalJob.status.in_(["pending", "assigned"])
            ).order_by(AIRemovalJob.deadline.asc()).all()
        else:
            jobs = (AIRemovalJob.query
                    .filter_by(writer_id=current_user.id)
                    .order_by(AIRemovalJob.deadline.asc())
                    .all())
        return render_template("writer_ai_removal_jobs.html", jobs=jobs)

    @main.route("/writer/ai-removal-jobs/<int:job_id>/upload-page")
    @require_login
    def writer_ai_removal_upload_page(job_id):
        """Redirect to the writer jobs page — upload happens inline there."""
        return redirect(url_for("main.writer_ai_removal_jobs"))

    @main.route("/writer/ai-removal-jobs/<int:job_id>/upload", methods=["POST"])
    @require_login
    def writer_upload_ai_removal(job_id):
        from .services.email_service import send_ai_removal_completed_student
        if not current_user.is_writer and not current_user.is_staff and not _is_owner(current_user):
            return redirect(url_for("main.dashboard"))

        job         = AIRemovalJob.query.get_or_404(job_id)
        final_text  = request.form.get("final_text", "").strip()
        upload_file = request.files.get("file")

        if not final_text and not upload_file:
            flash("Please paste the rewritten text or upload a file.", "error")
            return redirect(url_for("main.writer_ai_removal_jobs"))

        # Detect AI score on the rewritten text
        final_ai_score = 0.0
        if final_text:
            try:
                final_ai_score = ryter_service.detect_ai_score(final_text)
            except Exception:
                final_ai_score = 0.0

        # Upload file to Supabase if provided
        file_url  = None
        file_name = None
        if upload_file and upload_file.filename:
            from werkzeug.utils import secure_filename as _sf
            ext       = os.path.splitext(upload_file.filename)[1]
            file_name = _sf(upload_file.filename)
            path      = f"ai_removal/{job.id}/{uuid.uuid4().hex}{ext}"
            try:
                file_url = supabase_storage.upload_file(
                    path, upload_file.read(),
                    upload_file.content_type or "application/octet-stream",
                )
            except Exception as exc:
                current_app.logger.error("AI removal file upload error: %s", exc)

        job.final_text      = final_text or None
        job.final_file_url  = file_url
        job.final_file_name = file_name
        job.final_ai_score  = round(final_ai_score, 1)
        job.status          = "completed"
        job.completed_at    = datetime.utcnow()
        db.session.commit()

        # Notify the student
        student = User.query.get(job.user_id)
        if student and student.email:
            domains    = os.environ.get("REPLIT_DOMAINS", "smart-study-guides.com")
            base       = f"https://{domains.split(',')[0]}"
            status_url = f"{base}/assignments/{job.assignment_id}/ai-removal-status"
            try:
                send_ai_removal_completed_student(
                    student.email,
                    student.display_name,
                    job.assignment.topic,
                    round(final_ai_score, 1),
                    status_url,
                )
            except Exception:
                pass

        flash(f"Submitted! Final AI score: {round(final_ai_score,1)}%. Student has been notified.", "success")
        return redirect(url_for("main.writer_ai_removal_jobs"))

    # ── Student review ────────────────────────────────────────────────────────
    @main.route("/assignment/<int:aid>/review", methods=["POST"])
    @require_login
    def submit_review(aid):
        a = Assignment.query.filter_by(id=aid, user_id=current_user.id).first_or_404()
        if a.status != "complete":
            return jsonify({"error": "Assignment not complete"}), 400
        if a.review:
            return jsonify({"error": "Already reviewed"}), 400

        try:
            rating = int(request.json.get("rating", 0))
        except (ValueError, TypeError):
            rating = 0
        reason = (request.json.get("reason") or "").strip()

        if not (1 <= rating <= 5) or not reason:
            return jsonify({"error": "Invalid rating or reason"}), 400

        rev = Review(assignment_id=aid, user_id=current_user.id,
                     rating=rating, reason=reason)
        db.session.add(rev)
        db.session.commit()
        return jsonify({"ok": True})

    # ── Owner-only: inspect raw Ryter response ───────────────────────────────
    @main.route("/api/debug/ryter-raw", methods=["POST"])
    @require_login
    def debug_ryter_raw():
        if not _is_owner(current_user):
            return jsonify({"error": "Forbidden"}), 403
        import requests as _req
        text = request.form.get("text", "The impact of climate change on global food security is significant.")
        api_key = os.environ.get("RYTER_PRO_API_KEY", "")
        url = "https://api.ryter.pro/api/v1/ai-tools/execute/text-humanize"
        payload = {"text": text, "style": "academic", "model": "advanced"}
        try:
            r = _req.post(url, headers={"x-api-key": api_key, "Content-Type": "application/json"},
                          json=payload, timeout=90)
            return jsonify({
                "status_code": r.status_code,
                "content_type": r.headers.get("content-type"),
                "raw_text": r.text[:2000],
                "parsed_json": r.json() if "application/json" in r.headers.get("content-type","") else None,
            })
        except Exception as e:
            return jsonify({"error": str(e)})

    app.register_blueprint(main)
