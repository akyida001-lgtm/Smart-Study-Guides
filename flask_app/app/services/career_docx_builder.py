"""
Smart career document DOCX builder.
Detects job industry from position/doc_type and applies appropriate styling.
"""
import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Industry keyword detection ────────────────────────────────────────────────

_INDUSTRY_KEYWORDS = {
    "corporate": [
        "bank", "banking", "finance", "financial", "accounting", "accountant",
        "audit", "tax", "investment", "insurance", "legal", "law", "advocate",
        "compliance", "actuary", "actuarial", "treasury", "credit", "securities",
        "mortgage", "economist", "bursary",
    ],
    "creative": [
        "design", "designer", "graphic", "creative", "artist", "art ",
        "film", "fashion", "photography", "advertising", "marketing",
        "brand", "copywriter", "content creator", "illustrator", "ux",
        "ui ", "animation", "video", "media", "public relations", "pr ",
        "social media", "journalist", "editor", "storytelling",
    ],
    "technical": [
        "software", "developer", "engineer", "engineering", "data science",
        "data analyst", "it ", "tech", "programming", "programmer", "network",
        "cybersecurity", "machine learning", "artificial intelligence", "ai ",
        "devops", "cloud", "database", "systems analyst", "hardware",
        "embedded", "blockchain", "web developer", "full stack", "backend",
        "frontend", "mobile developer", "computer science",
    ],
    "hospitality": [
        "hotel", "hospitality", "catering", "food and", "restaurant",
        "chef", "tourism", "travel", "airline", "aviation", "front desk",
        "housekeeping", "concierge", "waiter", "barista", "bartender",
        "event", "events management", "wedding", "lodge", "resort",
    ],
    "medical": [
        "doctor", "nurse", "nursing", "medical", "clinical", "health",
        "pharmacy", "pharmacist", "dentist", "dental", "physiotherapy",
        "laboratory", "public health", "nutrition", "dietitian",
        "radiography", "optometry", "veterinary", "surgeon", "physician",
    ],
    "education": [
        "teacher", "teaching", "lecturer", "professor", "tutor",
        "education", "school", "academic", "trainer", "curriculum",
        "pedagogy", "early childhood", "special needs",
    ],
}


def _detect_industry(position: str, doc_label: str) -> str:
    text = (position + " " + doc_label).lower()
    for industry, keywords in _INDUSTRY_KEYWORDS.items():
        if any(k in text for k in keywords):
            return industry
    return "general"


# ── Style presets per industry ────────────────────────────────────────────────

_STYLES = {
    "corporate": {
        "font":        "Calibri",
        "title_size":  18, "heading_size": 13, "sub_size": 12, "body_size": 11,
        "heading_rgb": (0x1F, 0x38, 0x64),   # dark navy
        "divider_hex": "BFD3E8",
        "line_space":  1.15,
        "divider":     True,
    },
    "creative": {
        "font":        "Calibri",
        "title_size":  20, "heading_size": 13, "sub_size": 12, "body_size": 11,
        "heading_rgb": (0x5B, 0x21, 0xB6),   # purple
        "divider_hex": "DDD6FE",
        "line_space":  1.3,
        "divider":     True,
    },
    "technical": {
        "font":        "Arial",
        "title_size":  16, "heading_size": 12, "sub_size": 11, "body_size": 11,
        "heading_rgb": (0x1E, 0x40, 0xAF),   # blue
        "divider_hex": "BFDBFE",
        "line_space":  1.15,
        "divider":     False,
    },
    "hospitality": {
        "font":        "Calibri",
        "title_size":  18, "heading_size": 13, "sub_size": 12, "body_size": 11,
        "heading_rgb": (0x92, 0x40, 0x0E),   # warm amber-brown
        "divider_hex": "FDE68A",
        "line_space":  1.25,
        "divider":     True,
    },
    "medical": {
        "font":        "Calibri",
        "title_size":  18, "heading_size": 13, "sub_size": 12, "body_size": 11,
        "heading_rgb": (0x06, 0x5F, 0x46),   # dark green
        "divider_hex": "A7F3D0",
        "line_space":  1.2,
        "divider":     True,
    },
    "education": {
        "font":        "Calibri",
        "title_size":  18, "heading_size": 13, "sub_size": 12, "body_size": 11,
        "heading_rgb": (0x78, 0x35, 0x0F),   # auburn
        "divider_hex": "FED7AA",
        "line_space":  1.2,
        "divider":     True,
    },
    "general": {
        "font":        "Calibri",
        "title_size":  18, "heading_size": 13, "sub_size": 12, "body_size": 11,
        "heading_rgb": (0x1E, 0x3A, 0x5F),   # professional blue
        "divider_hex": "BFDBFE",
        "line_space":  1.2,
        "divider":     True,
    },
}


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _rgb(st: dict):
    r, g, b = st["heading_rgb"]
    return RGBColor(r, g, b)


def _set_run(run, font_name, size_pt, bold=False, color=None, italic=False):
    run.bold   = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def _set_para_spacing(p, line_space: float, space_before=0, space_after=4):
    fmt = p.paragraph_format
    fmt.line_spacing       = line_space
    fmt.space_before       = Pt(space_before)
    fmt.space_after        = Pt(space_after)


def _bottom_border(p, hex_color: str):
    """Add a thin bottom border to a paragraph (acts as section divider)."""
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _parse_inline(text: str):
    """Split text on **bold** markers → [(segment, is_bold), ...]"""
    parts = re.split(r"\*\*(.+?)\*\*", text)
    return [(part, idx % 2 == 1) for idx, part in enumerate(parts) if part]


def _add_runs(p, text: str, font: str, size: float, bold=False, color=None):
    """Add text to paragraph, honouring **bold** markers."""
    for seg, seg_bold in _parse_inline(text):
        run = p.add_run(seg)
        _set_run(run, font, size, bold=(bold or seg_bold), color=color)


# ── Section heading ───────────────────────────────────────────────────────────

def _add_heading(doc, text: str, st: dict, level: int = 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    size = st["heading_size"] if level == 1 else st["sub_size"]
    _add_runs(p, text, st["font"], size, bold=True, color=_rgb(st))
    _set_para_spacing(p, st["line_space"], space_before=8, space_after=3)
    if st["divider"] and level == 1:
        _bottom_border(p, st["divider_hex"])
    return p


# ── Title / name block ────────────────────────────────────────────────────────

def _add_title(doc, text: str, st: dict):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_run(run, st["font"], st["title_size"], bold=True, color=_rgb(st))
    _set_para_spacing(p, 1.0, space_before=0, space_after=2)
    return p


# ── Contact / metadata line ───────────────────────────────────────────────────

def _add_contact_line(doc, text: str, st: dict):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_run(run, st["font"], st["body_size"] - 0.5, color=RGBColor(0x64, 0x74, 0x8B))
    _set_para_spacing(p, 1.0, space_before=0, space_after=1)
    return p


# ── Body paragraph ────────────────────────────────────────────────────────────

def _add_body(doc, text: str, st: dict, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    _add_runs(p, text, st["font"], st["body_size"])
    _set_para_spacing(p, st["line_space"], space_before=0, space_after=4)
    return p


# ── Bullet point ─────────────────────────────────────────────────────────────

def _add_bullet(doc, text: str, st: dict):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_runs(p, text, st["font"], st["body_size"])
    _set_para_spacing(p, st["line_space"], space_before=0, space_after=2)
    return p


# ── Line classifier ───────────────────────────────────────────────────────────

_BULLET_RE   = re.compile(r"^[-•*◦·]\s+(.+)")
_NUMBERED_RE = re.compile(r"^\d+[.)]\s+(.+)")
_HASH1_RE    = re.compile(r"^#{1,2}\s*(.*)")


def _is_section_heading(line: str) -> bool:
    """ALL-CAPS line, 3-65 chars, not ending with a full stop."""
    s = line.strip()
    if not s or len(s) < 3 or len(s) > 65:
        return False
    if s.endswith(".") or s.endswith(","):
        return False
    upper = s.upper()
    return upper == s and any(c.isalpha() for c in s)


def _classify(line: str):
    """Return (kind, text) where kind ∈ h1/h2/bullet/body/blank."""
    s = line.strip()
    if not s:
        return "blank", ""
    m = _HASH1_RE.match(s)
    if m:
        level = len(re.match(r"^#+", s).group())
        return ("h1" if level == 1 else "h2"), m.group(1).strip()
    m = _BULLET_RE.match(s)
    if m:
        return "bullet", m.group(1).strip()
    m = _NUMBERED_RE.match(s)
    if m:
        return "bullet", m.group(0)
    if _is_section_heading(s):
        return "h1", s
    return "body", s


# ── Contact block heuristic ───────────────────────────────────────────────────

def _is_contact_line(line: str) -> bool:
    """True for lines that look like contact info (email, phone, URL, pipe-separated)."""
    s = line.strip()
    return bool(
        re.search(r"@[a-z0-9]", s, re.I)
        or re.search(r"\+?[\d][\d\s\-]{6,}", s)
        or re.search(r"https?://", s)
        or ("|" in s and len(s) < 120)
        or re.search(r"linkedin\.com", s, re.I)
    )


# ── Main builder ──────────────────────────────────────────────────────────────

def build_career_docx(content: str, doc_label: str, doc_type: str,
                      details: dict) -> bytes:
    position = details.get("position", "")
    industry = _detect_industry(position, doc_label)
    st       = _STYLES[industry]

    doc = Document()

    # Global Normal style
    normal = doc.styles["Normal"]
    normal.font.name = st["font"]
    normal.font.size = Pt(st["body_size"])

    # A4 margins (2.54 cm ≈ 1 inch)
    sec = doc.sections[0]
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Inches(1.0))

    lines = content.split("\n")
    n     = len(lines)
    i     = 0

    # ── Phase 1: consume a leading contact/header block ──────────────────────
    # Detect the first cluster of lines before the first heading or blank gap.
    # These get center-aligned special treatment (name + contact info).
    header_lines = []
    j = 0
    while j < n and j < 12:
        s = lines[j].strip()
        if not s:
            if header_lines:
                break        # blank after some header lines → end of block
            j += 1
            continue
        kind, _ = _classify(lines[j])
        if kind in ("h1", "h2") and header_lines:
            break            # a real heading after contact block → stop
        header_lines.append(lines[j].rstrip())
        j += 1

    if header_lines:
        first = header_lines[0].strip()
        # First non-empty line → title (name or document title)
        _add_title(doc, first, st)
        for hl in header_lines[1:]:
            s = hl.strip()
            if s:
                _add_contact_line(doc, s, st)
        # Separator space
        sp = doc.add_paragraph()
        _set_para_spacing(sp, 1.0, space_before=0, space_after=0)
        i = j
    else:
        i = 0

    # ── Phase 2: render the rest line by line ─────────────────────────────────
    blank_count = 0
    while i < n:
        raw  = lines[i]
        kind, text = _classify(raw)
        i += 1

        if kind == "blank":
            blank_count += 1
            if blank_count <= 1:
                sp = doc.add_paragraph()
                _set_para_spacing(sp, 1.0, space_before=0, space_after=0)
            continue

        blank_count = 0

        if kind == "h1":
            _add_heading(doc, text, st, level=1)
        elif kind == "h2":
            _add_heading(doc, text, st, level=2)
        elif kind == "bullet":
            _add_bullet(doc, text, st)
        else:
            # Detect address / date / salutation lines → align left, slightly indented
            _add_body(doc, text, st)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
